import os
import io
import tempfile
from datetime import date, datetime
from PIL import Image as PILImage

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
)
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.utils import ImageReader

from PySide6.QtWidgets import QFileDialog, QMessageBox

from paths import ICON_DIR
from utils.logo_manager import obtener_logo_bytes
from models.estu_model import MAPA_GRADOS
from models.institucion_model import InstitucionModel
from models.secciones_model import SeccionesModel
from models.anio_model import AnioEscolarModel
from models.emple_model import EmpleadoModel
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from utils.edad import calcular_edad
from utils.dialogs import crear_msgbox

# Estilos globales
styles = getSampleStyleSheet()

justificado = ParagraphStyle(
    name="Justificado",
    parent=styles["Normal"],
    alignment=TA_JUSTIFY,
    leading=18
)
centrado = ParagraphStyle(
    name="Centrado",
    parent=styles["Normal"],
    alignment=TA_CENTER,
    leading=18
)

page_width, page_height = letter


# UTILIDADES DE VALIDACIÓN Y CONVERSIÓN

def sanitizar_nombre_archivo(nombre: str) -> str:
    """Limpia caracteres inválidos de nombres de archivo."""
    caracteres_invalidos = r'<>:"/\|?*'
    for char in caracteres_invalidos:
        nombre = nombre.replace(char, '_')
    return nombre.strip()


def crear_carpeta_segura(ruta: str) -> tuple[bool, str]:
    """Crea una carpeta verificando permisos de escritura."""
    try:
        os.makedirs(ruta, exist_ok=True)
        
        # Verificar que sea escribible
        test_file = os.path.join(ruta, '.test_write')
        try:
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
            return True, "Carpeta creada correctamente"
        except Exception as e:
            return False, f"La carpeta no tiene permisos de escritura: {e}"
            
    except Exception as e:
        return False, f"No se pudo crear la carpeta: {e}"


def convertir_fecha_string(fecha) -> str:
    """Convierte diferentes formatos de fecha a string DD-MM-YYYY."""
    if fecha is None:
        return "N/A"
    
    if isinstance(fecha, (date, datetime)):
        return fecha.strftime("%d-%m-%Y")
    
    # Si ya es string, intentar parsearlo para validar
    if isinstance(fecha, str):
        for fmt in ("%d-%m-%Y", "%d/%m/%Y"):
            try:
                fecha_obj = datetime.strptime(fecha, fmt)
                return fecha_obj.strftime("%d-%m-%Y")
            except ValueError:
                continue
        return str(fecha)
    
    return str(fecha)

def formatear_anio(anio: int) -> str:
    """
    Formatea un año con punto de miles.
    """
    if not isinstance(anio, int):
        try:
            anio = int(anio)
        except (ValueError, TypeError):
            return str(anio)  # Retornar sin formatear si no es convertible
    
    # Formatear con separador de miles usando punto
    return f"{anio:,}".replace(',', '.')


def extraer_anio_escolar(anio_escolar: dict) -> tuple[int, int]:
    """Extrae años de inicio y fin del diccionario año_escolar."""
    if not anio_escolar:
        anio_actual = datetime.now().year
        return anio_actual, anio_actual + 1
    
    anio_inicio = anio_escolar.get('año_inicio')
    
    if isinstance(anio_inicio, (date, datetime)):
        anio_inicio = anio_inicio.year
    elif isinstance(anio_inicio, str):
        try:
            anio_inicio = int(anio_inicio.split('-')[0])
        except (ValueError, TypeError):
            anio_inicio = datetime.now().year
    else:
        anio_inicio = int(anio_inicio) if anio_inicio else datetime.now().year
    
    anio_fin = anio_inicio + 1
    return anio_inicio, anio_fin


def normalizar_cedula(cedula: str, es_estudiante: bool = False) -> str:
    """Normaliza formato de cédula, agregando prefijo V- o CE- si no tiene."""
    if not cedula:
        return "N/A"
    
    cedula = str(cedula).strip().upper()
    
    # Si ya tiene prefijo (V-, E-, J-, G-, CE-), mantenerlo
    if len(cedula) > 2:
        if cedula[:2] == 'CE' and cedula[2] == '-':
            return cedula
        if cedula[0] in ['V', 'E', 'J', 'G'] and cedula[1] == '-':
            return cedula
    
    # Si no tiene prefijo, agregar según el tipo
    prefijo = "CE-" if es_estudiante else "V-"
    return f"{prefijo}{cedula}"


def validar_datos_exportacion(datos: dict, campos_requeridos: list) -> tuple[bool, str]:
    """Valida que los datos tengan los campos requeridos para exportación."""
    if not datos:
        return False, "No se proporcionaron datos"
    
    campos_faltantes = []
    for campo in campos_requeridos:
        if campo not in datos or datos[campo] is None or str(datos[campo]).strip() == "":
            campos_faltantes.append(campo)
    
    if campos_faltantes:
        return False, f"Faltan campos requeridos: {', '.join(campos_faltantes)}"
    
    return True, "Datos válidos"


# FUNCIONES DE ENCABEZADO Y PIE DE PÁGINA

def draw_centered(canvas, text, y, font="Helvetica", size=12):
    """Dibuja texto centrado en el canvas."""
    canvas.setFont(font, size)
    text_width = canvas.stringWidth(text, font, size)
    x_center = (page_width - text_width) / 2
    canvas.drawString(x_center, y, text)


def encabezado(canvas, _doc, institucion_id=1):
    """
    Dibuja el encabezado con logo e institución en cada página del PDF.
    Prioriza el logo almacenado en la BD; si no existe, usa el archivo local.
    """
    logo_dibujado = False
    
    # Intentar cargar logo desde la base de datos
    try:
        logo_datos = obtener_logo_bytes()
        if logo_datos:
            logo_stream = io.BytesIO(logo_datos)
            logo_width = 65
            logo_height = 60
            x_center = (page_width - logo_width) / 2
            y_position = page_height - 180
            try:
                canvas.drawImage(
                    ImageReader(logo_stream),
                    x=x_center,
                    y=y_position,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=True,
                    mask='auto'
                )
                logo_dibujado = True
            except Exception as e:
                print(f"Error dibujando logo desde BD: {e}")
    except Exception:
        pass
    
    # Fallback: logo desde archivo local
    if not logo_dibujado:
        logo_path = os.path.join(ICON_DIR, "logo_escuela_fondo.png")
        if os.path.exists(logo_path):
            logo_width = 65
            logo_height = 60
            x_center = (page_width - logo_width) / 2
            y_position = page_height - 180
            try:
                canvas.drawImage(
                    logo_path,
                    x=x_center,
                    y=y_position,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=True
                )
            except Exception as e:
                print(f"Error dibujando logo desde archivo: {e}")
        else:
            print(f"Advertencia: Logo no encontrado en {logo_path}")

    # Obtener datos de institución
    institucion = InstitucionModel.obtener_por_id(institucion_id)

    if institucion:
        nombre = str(institucion.get("nombre", "")).upper()
        codigo = str(institucion.get("codigo_dea", ""))

        draw_centered(canvas, "REPÚBLICA BOLIVARIANA DE VENEZUELA", page_height - 50, "Helvetica", 10)
        draw_centered(canvas, "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN", page_height - 65, "Helvetica", 10)
        draw_centered(canvas, nombre, page_height - 80, "Helvetica", 10)
        draw_centered(canvas, f"CÓDIGO DEA: {codigo}", page_height - 95, "Helvetica", 10)
        draw_centered(canvas, "PUERTO LA CRUZ, EDO. ANZOÁTEGUI", page_height - 110, "Helvetica", 10)
    else:
        canvas.setFont("Helvetica-Bold", 12)
        fallback = "Institución no encontrada"
        text_width = canvas.stringWidth(fallback, "Helvetica-Bold", 12)
        x_center = (page_width - text_width) / 2
        canvas.drawString(x_center, page_height - 50, fallback)


def pie_pagina(canvas, _doc, institucion_id=1):
    """
    Dibuja el pie de página con datos de la institución.
    """
    institucion = InstitucionModel.obtener_por_id(institucion_id)

    font_name = "Helvetica"
    font_size = 10
    canvas.setFont(font_name, font_size)

    if institucion:
        direccion = str(institucion.get("direccion", "")).upper()
        telefono = str(institucion.get("telefono", ""))
        correo = str(institucion.get("correo", "")).upper()

        line1 = f"DIRECCIÓN: {direccion}" if direccion else ""
        line2 = f"TELÉFONO: {telefono} | CORREO: {correo}" if telefono or correo else ""
    else:
        line1 = ""
        line2 = ""

    for i, line in enumerate([line1, line2]):
        if line.strip():
            text_width = canvas.stringWidth(line, font_name, font_size)
            x_center = (page_width - text_width) / 2
            y_pos = 45 - i * (font_size + 2)
            canvas.drawString(x_center, y_pos, line)


def encabezado_y_pie(canvas, _doc):
    """
    Combina encabezado y pie de página.
    """
    encabezado(canvas, _doc)
    pie_pagina(canvas, _doc)


def _asegurar_png_compatible(ruta_imagen: str) -> str:
    """Convierte la imagen a PNG real si el formato no es compatible con python-docx."""
    try:
        with open(ruta_imagen, 'rb') as f:
            header = f.read(12)
        if header[:8] == b'\x89PNG\r\n\x1a\n':
            return ruta_imagen
    except Exception:
        return ruta_imagen

    try:
        img = PILImage.open(ruta_imagen)
        if img.mode == 'RGBA':
            fondo = PILImage.new('RGBA', img.size, (255, 255, 255, 0))
            fondo.paste(img, mask=img.split()[3] if 'A' in img.mode else None)
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        img.convert('RGBA').save(tmp.name, 'PNG')
        tmp.close()
        return tmp.name
    except Exception:
        return ruta_imagen


def encabezado_prosecucion(canvas, _doc):
    """Encabezado personalizado para constancia de prosecución."""
    # 1. Logo esquinado del ministerio (arriba a la izquierda)
    logo_ministerio_path = os.path.join(ICON_DIR, "ministerio.png")
    
    if os.path.exists(logo_ministerio_path):
        logo_min_width = 200
        logo_min_height = 120
        x_esquina = 40  # Margen izquierdo
        y_esquina = page_height - 100  # Posición superior

        try:
            canvas.drawImage(
                logo_ministerio_path,
                x_esquina,
                y_esquina,
                width=logo_min_width,
                height=logo_min_height,
                preserveAspectRatio=True,
                mask='auto'
            )
        except Exception as e:
            print(f"Error cargando logo del ministerio: {e}")
    
    # 2. Escudo centrado de Venezuela
    logo_venezuela_path = os.path.join(ICON_DIR, "Escudo_Venezuela.png")
        
    if os.path.exists(logo_venezuela_path):
        logo_ven_width = 70
        logo_ven_height = 85

        # Coordenadas centradas, debajo del logo del ministerio
        x_center = (page_width - logo_ven_width) / 2
        y_position = page_height - 170

        try:
            canvas.drawImage(
                logo_venezuela_path,
                x_center,
                y_position,
                width=logo_ven_width,
                height=logo_ven_height,
                preserveAspectRatio=True,
                mask='auto'
            )
        except Exception as e:
            print(f"Error cargando logo de Venezuela: {e}")


# FORMATOS ESTUDIANTES

def generar_constancia_estudios(estudiante: dict, institucion: dict) -> str:
    """Genera constancia de estudios en PDF para un estudiante."""
    # Validar datos requeridos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Grado", "Sección"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de estudios")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    # Nombre de archivo sanitizado
    nombre_base = sanitizar_nombre_archivo(f"Constancia_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=220,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE ESTUDIOS", styles["Title"]), Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        texto = (
            f"El suscrito, Director <b>PROF. {institucion['director'].upper()}</b>, portador de la Cédula de Identidad "
            f"<b>{director_ci}</b>, de la {institucion['nombre']}, hace constar que el(la) estudiante "
            f"<b>{estudiante['Apellidos']} {estudiante['Nombres']}</b>, "
            f"portador de la cédula escolar <b>{cedula_normalizada}</b>, cursa actualmente el "
            f"<b>{estudiante['Grado']} grado Sección '{estudiante['Sección']}'</b> de Educación Primaria en esta institución.<br/><br/>"
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>{dia}</b> días del mes de <b>{mes_es}</b> de <b>{anio}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 150))

        # Firma
        firma = f"________________________<br/>Prof. {institucion['director']}"
        story.append(Paragraph(firma, centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph(f"C.I. {director_ci}", centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph("Director", centrado))

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")


def generar_constancia_estudios_docx(estudiante: dict, institucion: dict) -> str:
    """Genera constancia de estudios en DOCX para un estudiante."""
    campos_est = ["Nombres", "Apellidos", "Cédula", "Grado", "Sección"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")

    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")

    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)
    director_ci = normalizar_cedula(institucion['director_ci'])

    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de estudios DOCX")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.docx")

    try:
        doc = Document()

        section = doc.sections[0]
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        logo_temp_path = None
        try:
            logo_datos = obtener_logo_bytes()
            if logo_datos:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    tmp.write(logo_datos)
                    logo_temp_path = tmp.name
        except Exception:
            pass

        if not logo_temp_path:
            logo_path_local = os.path.join(ICON_DIR, "logo_escuela_fondo.png")
            if os.path.exists(logo_path_local):
                logo_temp_path = logo_path_local

        if logo_temp_path and os.path.exists(logo_temp_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_temp_path, width=Inches(1.0))

        inst_data = InstitucionModel.obtener_por_id(1)
        nombre_inst = str(inst_data.get("nombre", "")).upper() if inst_data else ""
        codigo_inst = str(inst_data.get("codigo_dea", "")) if inst_data else ""

        lineas_membrete = [
            "REPÚBLICA BOLIVARIANA DE VENEZUELA",
            "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN",
            nombre_inst,
            f"CÓDIGO DEA: {codigo_inst}",
            "PUERTO LA CRUZ, EDO. ANZOÁTEGUI"
        ]
        for linea in lineas_membrete:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(linea)
            run.font.size = Pt(10)
            run.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_after = Pt(12)
        run_titulo = p_titulo.add_run("CONSTANCIA DE ESTUDIOS")
        run_titulo.bold = True
        run_titulo.font.size = Pt(16)
        run_titulo.font.name = "Arial"

        p_texto = doc.add_paragraph()
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.space_after = Pt(0)

        run1 = p_texto.add_run("El suscrito, Director ")
        run1.font.size = Pt(12)
        run1.font.name = "Arial"
        run2 = p_texto.add_run(f"PROF. {institucion['director'].upper()}")
        run2.bold = True
        run2.font.size = Pt(12)
        run2.font.name = "Arial"
        run3 = p_texto.add_run(", portador de la Cédula de Identidad ")
        run3.font.size = Pt(12)
        run3.font.name = "Arial"
        run4 = p_texto.add_run(director_ci)
        run4.bold = True
        run4.font.size = Pt(12)
        run4.font.name = "Arial"
        run5 = p_texto.add_run(f", de la {institucion['nombre']}, hace constar que el(la) estudiante ")
        run5.font.size = Pt(12)
        run5.font.name = "Arial"
        run6 = p_texto.add_run(f"{estudiante['Apellidos']} {estudiante['Nombres']}")
        run6.bold = True
        run6.font.size = Pt(12)
        run6.font.name = "Arial"
        run7 = p_texto.add_run(", portador de la cédula escolar ")
        run7.font.size = Pt(12)
        run7.font.name = "Arial"
        run8 = p_texto.add_run(cedula_normalizada)
        run8.bold = True
        run8.font.size = Pt(12)
        run8.font.name = "Arial"
        run9 = p_texto.add_run(f", cursa actualmente el ")
        run9.font.size = Pt(12)
        run9.font.name = "Arial"
        run10 = p_texto.add_run(f"{estudiante['Grado']} grado Sección '{estudiante['Sección']}'")
        run10.bold = True
        run10.font.size = Pt(12)
        run10.font.name = "Arial"
        run11 = p_texto.add_run(" de Educación Primaria en esta institución.")
        run11.font.size = Pt(12)
        run11.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(20)

        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year

        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fecha.paragraph_format.space_after = Pt(0)
        rf1 = p_fecha.add_run("Certificado que se expide en ")
        rf1.font.size = Pt(12)
        rf1.font.name = "Arial"
        rf2 = p_fecha.add_run("PUERTO LA CRUZ")
        rf2.bold = True
        rf2.font.size = Pt(12)
        rf2.font.name = "Arial"
        rf3 = p_fecha.add_run(", a los ")
        rf3.font.size = Pt(12)
        rf3.font.name = "Arial"
        rf4 = p_fecha.add_run(str(dia))
        rf4.bold = True
        rf4.font.size = Pt(12)
        rf4.font.name = "Arial"
        rf5 = p_fecha.add_run(" días del mes de ")
        rf5.font.size = Pt(12)
        rf5.font.name = "Arial"
        rf6 = p_fecha.add_run(mes_es)
        rf6.bold = True
        rf6.font.size = Pt(12)
        rf6.font.name = "Arial"
        rf7 = p_fecha.add_run(" de ")
        rf7.font.size = Pt(12)
        rf7.font.name = "Arial"
        rf8 = p_fecha.add_run(str(anio))
        rf8.bold = True
        rf8.font.size = Pt(12)
        rf8.font.name = "Arial"

        for _ in range(5):
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.space_before = Pt(0)

        p_linea = doc.add_paragraph()
        p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_linea.paragraph_format.space_after = Pt(3)
        run_linea = p_linea.add_run("________________________")
        run_linea.font.size = Pt(12)
        run_linea.font.name = "Arial"

        p_director = doc.add_paragraph()
        p_director.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_director.paragraph_format.space_after = Pt(3)
        run_dir = p_director.add_run(f"Prof. {institucion['director']}")
        run_dir.font.size = Pt(12)
        run_dir.font.name = "Arial"

        p_ci = doc.add_paragraph()
        p_ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ci.paragraph_format.space_after = Pt(3)
        run_ci = p_ci.add_run(f"C.I. {director_ci}")
        run_ci.font.size = Pt(12)
        run_ci.font.name = "Arial"

        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cargo = p_cargo.add_run("Director")
        run_cargo.font.size = Pt(12)
        run_cargo.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        if inst_data:
            direccion = str(inst_data.get("direccion", "")).upper()
            telefono = str(inst_data.get("telefono", ""))
            correo = str(inst_data.get("correo", "")).upper()

            line1 = f"DIRECCIÓN: {direccion}" if direccion else ""
            line2 = f"TELÉFONO: {telefono} | CORREO: {correo}" if telefono or correo else ""

            for linea_pie in [line1, line2]:
                if linea_pie.strip():
                    p_pie = doc.add_paragraph()
                    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_pie.paragraph_format.space_after = Pt(1)
                    run_pie = p_pie.add_run(linea_pie)
                    run_pie.font.size = Pt(10)
                    run_pie.font.name = "Arial"

        doc.save(nombre_archivo)
        return nombre_archivo

    except Exception as e:
        raise IOError(f"Error generando DOCX: {e}")
    finally:
        if logo_temp_path and logo_temp_path != os.path.join(ICON_DIR, "logo_escuela_fondo.png"):
            try:
                os.unlink(logo_temp_path)
            except OSError:
                pass


def generar_certificado_promocion_sexto_docx(estudiante: dict, institucion: dict, anio_escolar_egreso: str) -> str:
    """Genera certificado de promoción de 6to grado en DOCX."""
    estudiante_norm = {"Nombres": estudiante.get("Nombres") or estudiante.get("nombres", ""),
                       "Apellidos": estudiante.get("Apellidos") or estudiante.get("apellidos", ""),
                       "Cédula": estudiante.get("Cédula") or estudiante.get("cedula", ""),
                       "Ciudad": estudiante.get("Ciudad") or estudiante.get("ciudad", ""),
                       "Fecha Nac.": estudiante.get("Fecha Nac.") or estudiante.get("fecha_nac")}

    campos_est = ["Nombres", "Apellidos", "Cédula", "Ciudad", "Fecha Nac."]
    valido, mensaje = validar_datos_exportacion(estudiante_norm, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")

    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")

    if not anio_escolar_egreso:
        raise ValueError("Año escolar de egreso no proporcionado")

    estudiante_norm["Nombres"] = str(estudiante_norm["Nombres"]).strip().upper()
    estudiante_norm["Apellidos"] = str(estudiante_norm["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante_norm["Cédula"], es_estudiante=True)
    fecha_nac_str = convertir_fecha_string(estudiante_norm['Fecha Nac.'])

    anio_periodo = str(anio_escolar_egreso).replace('/', '-')
    anios = anio_periodo.split('-')
    if len(anios) == 2:
        try:
            anio_inicio_egreso = int(anios[0].strip())
            anio_fin_egreso = int(anios[1].strip())
            anio_periodo_form = f"{formatear_anio(anio_inicio_egreso)}-{formatear_anio(anio_fin_egreso)}"
        except ValueError:
            anio_periodo_form = anio_periodo
    else:
        anio_periodo_form = anio_periodo

    carpeta = os.path.join(os.getcwd(), "exportados", "Certificados de promocion DOCX")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Certificado_prosecusion_primaria{estudiante_norm['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.docx")

    temp_files = []

    try:
        doc = Document()

        section = doc.sections[0]
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        logo_ministerio_path = None
        logo_venezuela_path = None
        logo_ministerio_png = None
        logo_venezuela_png = None

        logo_ministerio_local = os.path.join(ICON_DIR, "ministerio.png")
        if os.path.exists(logo_ministerio_local):
            logo_ministerio_path = logo_ministerio_local
            logo_ministerio_png = _asegurar_png_compatible(logo_ministerio_local)
            if logo_ministerio_png != logo_ministerio_local:
                temp_files.append(logo_ministerio_png)

        logo_venezuela_local = os.path.join(ICON_DIR, "Escudo_Venezuela.png")
        if os.path.exists(logo_venezuela_local):
            logo_venezuela_path = logo_venezuela_local
            logo_venezuela_png = _asegurar_png_compatible(logo_venezuela_local)
            if logo_venezuela_png != logo_venezuela_local:
                temp_files.append(logo_venezuela_png)

        if logo_ministerio_png:
            p_logo_min = doc.add_paragraph()
            p_logo_min.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_logo_min = p_logo_min.add_run()
            run_logo_min.add_picture(logo_ministerio_png, width=Inches(1.8))

        if logo_venezuela_png:
            p_logo_ven = doc.add_paragraph()
            p_logo_ven.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo_ven = p_logo_ven.add_run()
            run_logo_ven.add_picture(logo_venezuela_png, width=Inches(0.65))

        inst_data = InstitucionModel.obtener_por_id(1)
        nombre_inst = str(inst_data.get("nombre", "")).upper() if inst_data else ""
        codigo_inst = str(inst_data.get("codigo_dea", "")) if inst_data else ""

        lineas_membrete = [
            "REPÚBLICA BOLIVARIANA DE VENEZUELA",
            "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN",
            nombre_inst,
            f"CÓDIGO DEA: {codigo_inst}",
            "PUERTO LA CRUZ, EDO. ANZOÁTEGUI"
        ]
        for linea in lineas_membrete:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(linea)
            run.font.size = Pt(10)
            run.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_after = Pt(12)
        run_titulo = p_titulo.add_run("CERTIFICADO DE EDUCACIÓN PRIMARIA")
        run_titulo.bold = True
        run_titulo.font.size = Pt(16)
        run_titulo.font.name = "Arial"

        director_ci = normalizar_cedula(institucion['director_ci'])
        ultima_seccion = (estudiante.get('ultima_seccion') or
                          estudiante.get('Ultima Seccion') or
                          estudiante.get('letra') or 'N/A')

        p_texto = doc.add_paragraph()
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.space_after = Pt(0)

        runs_data = [
            ("Quien suscribe ", False),
            (institucion['director'].upper(), True),
            (" titular de la Cédula de Identidad Nº ", False),
            (director_ci, True),
            (" Director(a) de la Institución Educativa ", False),
            (institucion['nombre'].upper(), True),
            (", ubicada en el Municipio ", False),
            ("JUAN ANTONIO SOTILLO", True),
            (" de la Parroquia ", False),
            ("PUERTO LA CRUZ", True),
            (" adscrita al Centro de Desarrollo de la Calidad Educativa Estadal ", False),
            ("ANZOÁTEGUI", True),
            (". Por la presente certifica que el (la) estudiante ", False),
            (f"{estudiante_norm['Apellidos']} {estudiante_norm['Nombres']}", True),
            (" titular de Cédula Escolar Nº ", False),
            (cedula_normalizada, True),
            (", nacido (a) en el Municipio ", False),
            (estudiante_norm['Ciudad'], True),
            (" del Estado ", False),
            ("ANZOÁTEGUI", True),
            (", en fecha ", False),
            (fecha_nac_str, True),
            (", cursó el ", False),
            ("6to Grado", True),
            (" correspondiéndole el literal ", False),
            (ultima_seccion, True),
            (" durante el periodo escolar ", False),
            (anio_periodo_form, True),
            (", siendo promovido(a) al ", False),
            ("1er Año del Nivel de Educación Media", True),
            (", previo cumplimiento a los requisitos establecidos en la normativa legal vigente.", False),
        ]

        for texto_run, is_bold in runs_data:
            run = p_texto.add_run(texto_run)
            run.bold = is_bold
            run.font.size = Pt(12)
            run.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(20)

        fecha_hoy = date.today()
        anio_actual = fecha_hoy.year
        anio_actual_form = formatear_anio(anio_actual)

        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fecha.paragraph_format.space_after = Pt(0)

        fecha_runs = [
            ("Certificado que se expide en ", False),
            ("PUERTO LA CRUZ", True),
            (", a los ", False),
            ("30", True),
            (" días del mes de ", False),
            ("Julio", True),
            (" de ", False),
            (anio_actual_form, True),
        ]
        for texto_run, is_bold in fecha_runs:
            run = p_fecha.add_run(texto_run)
            run.bold = is_bold
            run.font.size = Pt(12)
            run.font.name = "Arial"

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        tabla_firmas = doc.add_table(rows=5, cols=2)
        tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla_firmas.style = 'Table Grid'

        def _set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Arial"

        def _set_cell_multi_line(cell, lines_data, alignment=WD_ALIGN_PARAGRAPH.CENTER):
            cell.text = ""
            for i, (text, bold, size) in enumerate(lines_data):
                if i == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.alignment = alignment
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(size)
                run.font.name = "Arial"

        _set_cell_multi_line(tabla_firmas.rows[0].cells[0], [
            ("INSTITUCIÓN EDUCATIVA", True, 9),
            ("(PARA VALIDEZ NACIONAL)", True, 9),
        ])
        _set_cell_multi_line(tabla_firmas.rows[0].cells[1], [
            ("CENTRO DE DESARROLLO DE LA CALIDAD", True, 9),
            ("EDUCATIVA ESTADAL", True, 9),
            ("(PARA VALIDEZ INTERNACIONAL)", True, 9),
        ])

        _set_cell_text(tabla_firmas.rows[1].cells[0], "DIRECTOR(A)", bold=True, size=10)
        _set_cell_text(tabla_firmas.rows[1].cells[1], "DIRECTOR(A)", bold=True, size=10)

        _set_cell_text(tabla_firmas.rows[2].cells[0],
                       f"Nombre y Apellido: {institucion['director'].upper()}",
                       bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(tabla_firmas.rows[2].cells[1],
                       "Nombre y Apellido:",
                       bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        _set_cell_text(tabla_firmas.rows[3].cells[0],
                       f"Número de C.I: {director_ci}",
                       bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(tabla_firmas.rows[3].cells[1],
                       "Número de C.I:",
                       bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        _set_cell_multi_line(tabla_firmas.rows[4].cells[0], [
            ("Firma y Sello:", False, 10),
            ("", False, 10),
            ("", False, 10),
        ], alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_multi_line(tabla_firmas.rows[4].cells[1], [
            ("Firma y Sello:", False, 10),
            ("", False, 10),
            ("", False, 10),
        ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

        doc.save(nombre_archivo)
        return nombre_archivo

    except Exception as e:
        raise IOError(f"Error generando DOCX: {e}")
    finally:
        for tmp in temp_files:
            try:
                os.unlink(tmp)
            except OSError:
                pass


def generar_buena_conducta(estudiante: dict, institucion: dict, anio_escolar: dict) -> str:
    """Genera constancia de buena conducta en PDF para un estudiante."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Grado", "Sección"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    # Extraer años
    anio_inicio, anio_fin = extraer_anio_escolar(anio_escolar)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)

    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de buena conducta")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_buena_conducta_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=220,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE BUENA CONDUCTA", styles["Title"]), Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        texto = (
            f"El suscrito, Director <b>PROF. {institucion['director'].upper()}</b>, portador de la Cédula de Identidad "
            f"<b>{director_ci}</b>, de la {institucion['nombre']}, que funciona en Puerto La Cruz, "
            f"hace constar que el alumno(a): <b>{estudiante['Apellidos']} {estudiante['Nombres']}</b>, "
            f"portador de la cédula de identidad <b>{cedula_normalizada}</b>, estudiante del "
            f"<b>{estudiante['Grado']} Grado Sección '{estudiante['Sección']}'</b> "
            f"de Educación Primaria durante el Año Escolar <b>{anio_inicio}-{anio_fin}</b>, mantuvo una "
            f"<b>Buena Conducta</b> durante su permanencia en esta institución educativa.<br/><br/>"
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>{dia}</b> días del mes de <b>{mes_es}</b> de <b>{anio}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 150))

        # Firma
        firma = f"________________________<br/>Prof. {institucion['director']}"
        story.append(Paragraph(firma, centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph(f"C.I. {director_ci}", centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph("Director", centrado))

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")


def generar_constancia_inscripcion(estudiante: dict, institucion: dict) -> str:
    """Genera constancia de inscripción en PDF para un estudiante."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Grado", "Ciudad", "Fecha Nac.", "Fecha Ingreso"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    estudiante["Cédula"] = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)

    # Convertir fechas
    fecha_nac_str = convertir_fecha_string(estudiante['Fecha Nac.'])
    fecha_ingreso_str = convertir_fecha_string(estudiante['Fecha Ingreso'])
    
    # Calcular edad
    try:
        fecha_nac = estudiante['Fecha Nac.']
        if isinstance(fecha_nac, (date, datetime)):
            edad = calcular_edad(fecha_nac)
        else:
            try:
                fecha_obj = datetime.strptime(fecha_nac_str, "%d-%m-%Y").date()
                edad = calcular_edad(fecha_obj)
            except (ValueError, TypeError):
                edad = "N/A"
    except (ValueError, TypeError, KeyError):
        edad = "N/A"

    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de inscripcion")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_inscripcion_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=220,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE INSCRIPCIÓN", styles["Title"]), Spacer(1, 16)]

        # Texto principal
        texto = (
            f"La Dirección del plantel hace constar mediante la presente, que el Alumno(a): "
            f"<b>{estudiante['Apellidos']} {estudiante['Nombres']}</b>, nacido en {estudiante['Ciudad']} "
            f"en fecha <b>{fecha_nac_str}</b>, de <b>{edad}</b> años de edad, fué inscrito "
            f"en esta institución el día <b>{fecha_ingreso_str}</b> para cursar el <b>{estudiante['Grado']} Grado</b> "
            f"de Educación Primaria."
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>{dia}</b> días del mes de <b>{mes_es}</b> de <b>{anio}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 150))

        # Firma
        director_ci = normalizar_cedula(institucion['director_ci'])
        firma = f"________________________<br/>Prof. {institucion['director']}"
        story.append(Paragraph(firma, centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph(f"C.I. {director_ci}", centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph("Director", centrado))

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")


def generar_constancia_prosecucion_inicial(estudiante: dict, institucion: dict, anio_escolar: dict) -> str:
    """Genera constancia de prosecución de educación inicial a primaria en PDF."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Ciudad", "Fecha Nac."]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    anio_inicio, anio_fin = extraer_anio_escolar(anio_escolar)
    
    # Formatear años con punto de miles
    anio_inicio_form = formatear_anio(anio_inicio)
    anio_fin_form = formatear_anio(anio_fin)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)
    
    # Convertir fecha
    fecha_nac_str = convertir_fecha_string(estudiante['Fecha Nac.'])
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de prosecución inicial")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_prosecucion_inicial_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=180,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE PROSECUSIÓN<br/>EN EL NIVEL DE EDUCACIÓN INICIAL", styles["Title"]),
                 Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        texto = (
            f"Quien suscribe <b>{institucion['director'].upper()}</b> titular de la Cédula de Identidad "
            f"Nº <b>{director_ci}</b> Director(a) de la Institución Educativa "
            f"<b>{institucion['nombre'].upper()}</b>, ubicada en el Municipio <b>JUAN ANTONIO SOTILLO</b> de la Parroquia "
            f"<b>PUERTO LA CRUZ</b> adscrita al Centro de Desarrollo de la Calidad Educativa Estadal <b>ANZOÁTEGUI</b>. "
            f"Por la presente certifica que el (la) estudiante <b>{estudiante['Apellidos']} {estudiante['Nombres']}</b> "
            f"titular de Cédula Escolar <b>{cedula_normalizada}</b>, nacido (a) en el municipio <b>{estudiante['Ciudad']}</b> "
            f"del Estado <b>ANZOÁTEGUI</b> en fecha <b>{fecha_nac_str}</b>, cursó el <b>3er nivel</b> de la etapa de "
            f"<b>Educación Inicial</b> durante el periodo escolar <b>{anio_inicio_form}-{anio_fin_form}</b>, siendo "
            f"promovido(a) a <b>primer grado de primaria</b>, previo cumplimiento a los requisitos establecidos "
            f"en la normativa legal vigente."
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        anio_actual = fecha_hoy.year
        anio_actual_form = formatear_anio(anio_actual)
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>30</b> días del mes de <b>Julio</b> de <b>{anio_actual_form}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 30))

        # Tabla de firmas institucionales
        tabla_firmas_data = [
            # Encabezados
            [
                Paragraph("<b>INSTITUCIÓN EDUCATIVA<br/>(PARA VALIDEZ NACIONAL)</b>", centrado),
                Paragraph("<b>CENTRO DE DESARROLLO DE LA CALIDAD<br/>EDUCATIVA ESTADAL<br/>(PARA VALIDEZ INTERNACIONAL)</b>", centrado)
            ],
            # Director(a)
            [
                Paragraph("DIRECTOR(A)", styles["Normal"]),
                Paragraph("DIRECTOR(A)", styles["Normal"])
            ],
            # Nombre y Apellido
            [
                Paragraph(f"Nombre y Apellido: {institucion['director'].upper()}", styles["Normal"]),
                Paragraph("Nombre y Apellido:", styles["Normal"])
            ],
            # Cédula de Identidad
            [
                Paragraph(f"Número de C.I: {director_ci}", styles["Normal"]),
                Paragraph("Número de C.I:", styles["Normal"])
            ],
            # Firma y Sello (espacios vacíos)
            [
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"]),
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"])
            ]
        ]

        tabla_firmas = Table(tabla_firmas_data, colWidths=[page_width/2 - 100, page_width/2 - 100])
        tabla_firmas.setStyle(TableStyle([
            # Bordes externos de la tabla
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            
            # Alineación
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Encabezados centrados
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),   # Resto alineado a la izquierda
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),   # Alineación vertical superior
            
            # Padding interno
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(tabla_firmas)

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_prosecucion, onLaterPages=encabezado_prosecucion)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")

def generar_certificado_prosecucion_primaria(estudiante: dict, institucion: dict, anio_escolar: dict) -> str:
    """Genera constancia de prosecución de primaria en PDF."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Ciudad", "Fecha Nac."]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    anio_inicio, anio_fin = extraer_anio_escolar(anio_escolar)
    
    # Formatear años con punto de miles
    anio_inicio_form = formatear_anio(anio_inicio)
    anio_fin_form = formatear_anio(anio_fin)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)
    
    # Convertir fecha
    fecha_nac_str = convertir_fecha_string(estudiante['Fecha Nac.'])
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de prosecución primaria")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_prosecucion_primaria_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=180,
            bottomMargin=50
        )

        story = [Paragraph("CERTIFICADO<br/>DE PROSECUCIÓN PRIMARIA", styles["Title"]),
                 Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        texto = (
            f"Quien suscribe <b>{institucion['director'].upper()}</b> titular de la Cédula de Identidad "
            f"Nº <b>{director_ci}</b> Director(a) de la Institución Educativa "
            f"<b>{institucion['nombre'].upper()}</b>, ubicada en el Municipio <b>JUAN ANTONIO SOTILLO</b> de la Parroquia "
            f"<b>PUERTO LA CRUZ</b> adscrita al Centro de Desarrollo de la Calidad Educativa Estadal <b>ANZOÁTEGUI</b>. "
            f"Por la presente certifica que el (la) estudiante <b>{estudiante['Apellidos']} {estudiante['Nombres']}</b> "
            f"titular de Cédula Escolar <b>{cedula_normalizada}</b>, nacido (a) en el municipio <b>{estudiante['Ciudad']}</b> "
            f"del Estado <b>ANZOÁTEGUI</b> en fecha <b>{fecha_nac_str}</b>, cursó el <b>{estudiante["Grado"]} Grado</b> correspondiéndole "
            f"el literal <b>{estudiante["Sección"]}</b> durante el periodo escolar <b>{anio_inicio_form}-{anio_fin_form}</b>, <b>siendo "
            f"promovido(a) al {MAPA_GRADOS.get(estudiante["Grado"], "Grado desconocido.")} Grado de Educación Primaria</b>, previo "
            f"cumplimiento a los requisitos establecidos en la normativa legal vigente."
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>{dia}</b> días del mes de <b>{mes_es}</b> de <b>{anio}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 30))

        # Tabla de firmas institucionales
        tabla_firmas_data = [
            # Encabezados
            [
                Paragraph("<b>INSTITUCIÓN EDUCATIVA<br/>(PARA VALIDEZ NACIONAL)</b>", centrado),
                Paragraph("<b>CENTRO DE DESARROLLO DE LA CALIDAD<br/>EDUCATIVA ESTADAL<br/>(PARA VALIDEZ INTERNACIONAL)</b>", centrado)
            ],
            # Director(a)
            [
                Paragraph("DIRECTOR(A)", styles["Normal"]),
                Paragraph("DIRECTOR(A)", styles["Normal"])
            ],
            # Nombre y Apellido
            [
                Paragraph(f"Nombre y Apellido: {institucion['director'].upper()}", styles["Normal"]),
                Paragraph("Nombre y Apellido:", styles["Normal"])
            ],
            # Cédula de Identidad
            [
                Paragraph(f"Número de C.I: {director_ci}", styles["Normal"]),
                Paragraph("Número de C.I:", styles["Normal"])
            ],
            # Firma y Sello (espacios vacíos)
            [
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"]),
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"])
            ]
        ]

        tabla_firmas = Table(tabla_firmas_data, colWidths=[page_width/2 - 100, page_width/2 - 100])
        tabla_firmas.setStyle(TableStyle([
            # Bordes externos de la tabla
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            
            # Alineación
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Encabezados centrados
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),   # Resto alineado a la izquierda
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),   # Alineación vertical superior
            
            # Padding interno
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(tabla_firmas)

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_prosecucion, onLaterPages=encabezado_prosecucion)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")

# FORMATOS EMPLEADOS

def generar_constancia_trabajo(empleado: dict, institucion: dict) -> str:
    """Genera constancia de trabajo en PDF para un empleado."""
    # Validar datos
    campos_emp = ["Nombres", "Apellidos", "Cédula", "Cargo", "Fecha Ingreso"]
    valido, mensaje = validar_datos_exportacion(empleado, campos_emp)
    if not valido:
        raise ValueError(f"Datos de empleado incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    # Normalizar datos
    empleado["Nombres"] = str(empleado["Nombres"]).strip().upper()
    empleado["Apellidos"] = str(empleado["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(empleado["Cédula"])
    
    # Convertir fecha
    fecha_ingreso_str = convertir_fecha_string(empleado['Fecha Ingreso'])
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de trabajo")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"ConstanciaTrabajo_{empleado['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=220,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE TRABAJO", styles["Title"]), Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        texto = (
            f"Quien suscribe, <b>{institucion['director']}</b> Cédula de identidad <b>{director_ci}</b> "
            f"Director(a) de la {institucion['nombre']} hace constar "
            f"por medio de la presente que el (la) ciudadano (a) <b>{empleado['Nombres']} {empleado['Apellidos']}</b>, Cédula de Identidad "
            f"<b>{cedula_normalizada}</b> presta su servicio como <b>{empleado['Cargo']}</b> en esta institución desde "
            f"el <b>{fecha_ingreso_str}</b> hasta la presente fecha."
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = f"Certificado que se expide a petición de la parte interesada, en <b>PUERTO LA CRUZ</b>, a los <b>{dia}</b> días del mes de <b>{mes_es}</b> de <b>{anio}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 150))

        # Firma
        firma = f"________________________<br/>Prof. {institucion['director']}"
        story.append(Paragraph(firma, centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph(f"C.I. {director_ci}", centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph("Director", centrado))

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")


# REPORTES Y EXPORTACIONES

def exportar_reporte_pdf(parent, figure, titulo, criterio, etiquetas, valores, total) -> str:
    """Exporta un reporte estadístico a PDF con gráfica y tabla."""
    try:
        sugerido = f"reporte_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        ruta, _ = QFileDialog.getSaveFileName(
            parent,
            "Guardar reporte",
            sugerido,
            "PDF Files (*.pdf)"
        )
        
        if not ruta:
            return None
            
        if not ruta.endswith(".pdf"):
            ruta += ".pdf"

        # Guardar la figura en memoria
        buf = io.BytesIO()
        figure.savefig(buf, format="png", bbox_inches="tight", dpi=150)
        buf.seek(0)

        # Documento
        doc = SimpleDocTemplate(
            ruta,
            pagesize=letter,
            leftMargin=70,
            rightMargin=70,
            topMargin=180,
            bottomMargin=60
        )
        
        story = []

        # Título y fecha
        fecha = datetime.now().strftime("%d-%m-%Y %H:%M")
        titulo_style = ParagraphStyle(
            'TituloReporte',
            parent=styles['Title'],
            fontSize=16,
            textColor=colors.HexColor("#2C3E50"),
            spaceAfter=8
        )
        story.append(Paragraph(f"<b>{titulo}</b>", titulo_style))
        story.append(Spacer(1, 5))
        
        # Subtítulo con fecha
        subtitulo_style = ParagraphStyle(
            'Subtitulo',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor("#7F8C8D"),
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generado el {fecha}", subtitulo_style))
        story.append(Spacer(1, 20))

        # Calcular estadísticas
        if valores:
            promedio = sum(valores) / len(valores)
            maximo = max(valores)
            minimo = min(valores)
            
            # Encontrar categoría con valor máximo
            idx_max = valores.index(maximo)
            categoria_max = etiquetas[idx_max]
            
            # Calcular porcentaje del máximo
            porcentaje_max = (maximo / total * 100) if total > 0 else 0
        else:
            promedio = maximo = minimo = 0
            categoria_max = "N/A"
            porcentaje_max = 0

        # Crear tabla de resumen
        resumen_data = [
            ["RESUMEN", ""],
            ["Total de registros", f"{total}"],
            ["Promedio por categoría", f"{promedio:.1f}"],
            ["Valor máximo", f"{maximo} ({categoria_max})"],
            ["Porcentaje máximo", f"{porcentaje_max:.1f}%"],
            ["Categorías analizadas", f"{len(etiquetas)}"]
        ]
        
        resumen_table = Table(resumen_data, colWidths=[180, 180])
        resumen_table.setStyle(TableStyle([
            # Encabezado
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#3498DB")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,0), 11),
            ("ALIGN", (0,0), (-1,0), "CENTER"),
            ("SPAN", (0,0), (-1,0)),
            
            # Contenido
            ("BACKGROUND", (0,1), (0,-1), colors.HexColor("#ECF0F1")),
            ("BACKGROUND", (1,1), (1,-1), colors.white),
            ("FONTNAME", (0,1), (0,-1), "Helvetica-Bold"),
            ("FONTNAME", (1,1), (1,-1), "Helvetica"),
            ("FONTSIZE", (0,1), (-1,-1), 9),
            ("ALIGN", (0,1), (0,-1), "LEFT"),
            ("ALIGN", (1,1), (1,-1), "RIGHT"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            
            # Bordes
            ("BOX", (0,0), (-1,-1), 1.5, colors.HexColor("#3498DB")),
            ("LINEBELOW", (0,0), (-1,0), 1.5, colors.HexColor("#2980B9")),
            ("INNERGRID", (0,1), (-1,-1), 0.5, colors.HexColor("#BDC3C7")),
            
            # Padding
            ("LEFTPADDING", (0,0), (-1,-1), 10),
            ("RIGHTPADDING", (0,0), (-1,-1), 10),
            ("TOPPADDING", (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(resumen_table)
        story.append(Spacer(1, 25))

        # Preparar datos con porcentajes
        data = [["Categoría", "Cantidad", "Porcentaje"]]
        for e, v in zip(etiquetas, valores):
            porcentaje = (v / total * 100) if total > 0 else 0
            data.append([str(e), str(v), f"{porcentaje:.1f}%"])
        
        # Fila de total
        data.append(["TOTAL", str(total), "100%"])

        # Crear tabla
        table = Table(data, hAlign="CENTER", colWidths=[190, 90, 80])
        
        # Estilo degradado para filas alternas
        table_style = [
            # Encabezado
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#34495E")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,0), 10),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            
            # Contenido
            ("FONTNAME", (0,1), (-1,-2), "Helvetica"),
            ("FONTSIZE", (0,1), (-1,-2), 9),
            
            # Fila total
            ("BACKGROUND", (0,-1), (-1,-1), colors.HexColor("#E8F6F3")),
            ("TEXTCOLOR", (0,-1), (-1,-1), colors.HexColor("#117A65")),
            ("FONTNAME", (0,-1), (-1,-1), "Helvetica-Bold"),
            ("FONTSIZE", (0,-1), (-1,-1), 10),
            
            # Bordes
            ("BOX", (0,0), (-1,-1), 1.5, colors.HexColor("#34495E")),
            ("LINEBELOW", (0,0), (-1,0), 1.5, colors.HexColor("#2C3E50")),
            ("INNERGRID", (0,0), (-1,-1), 0.5, colors.HexColor("#BDC3C7")),
            
            # Padding
            ("LEFTPADDING", (0,0), (-1,-1), 8),
            ("RIGHTPADDING", (0,0), (-1,-1), 8),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]
        
        # Agregar colores alternados para filas (excepto encabezado y total)
        for i in range(1, len(data) - 1):
            if i % 2 == 0:
                table_style.append(("BACKGROUND", (0,i), (-1,i), colors.HexColor("#F8F9F9")))
            else:
                table_style.append(("BACKGROUND", (0,i), (-1,i), colors.white))
        
        table.setStyle(TableStyle(table_style))
        story.append(table)
        story.append(Spacer(1, 25))

        story.append(Paragraph("<b>Representación Gráfica</b>", styles["Heading2"]))
        story.append(Spacer(1, 10))
        story.append(Image(buf, width=450, height=330))
        story.append(Spacer(1, 15))

        # ============== INTERPRETACIÓN AUTOMÁTICA ==============
        interpretacion_style = ParagraphStyle(
            'Interpretacion',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor("#34495E"),
            alignment=TA_JUSTIFY,
            leading=12,
            leftIndent=20,
            rightIndent=20,
            spaceBefore=5,
            spaceAfter=5
        )
        
        interpretacion = (
            f"<b>Análisis:</b> Se analizaron <b>{len(etiquetas)}</b> categorías según el criterio "
            f"<i>{criterio}</i>, con un total de <b>{total}</b> registros. "
            f"La categoría con mayor representación es <b>{categoria_max}</b> con <b>{maximo}</b> "
            f"registros ({porcentaje_max:.1f}% del total)."
        )
        
        if len(valores) > 1:
            # Calcular desviación entre máximo y mínimo
            rango = maximo - minimo
            if total > 0:
                variabilidad = (rango / total * 100)
                if variabilidad > 50:
                    interpretacion += " Se observa una <b>alta variabilidad</b> en la distribución."
                elif variabilidad > 20:
                    interpretacion += " La distribución muestra <b>variabilidad moderada</b>."
                else:
                    interpretacion += " La distribución es relativamente <b>homogénea</b>."
        
        story.append(Paragraph(interpretacion, interpretacion_style))
        story.append(Spacer(1, 20))

        # ============== OBSERVACIONES ==============
        story.append(Paragraph("<b>Observaciones Adicionales:</b>", styles["Normal"]))
        story.append(Spacer(1, 8))
        for _ in range(3):
            story.append(Paragraph("_" * 90, styles["Normal"]))
            story.append(Spacer(1, 10))

        # ============== FOOTER PERSONALIZADO ==============
        def footer_metadata(canvas, _doc):
            """Footer con metadata del sistema"""
            canvas.saveState()
            
            # Primero dibujamos el encabezado y pie institucional
            encabezado(canvas, _doc)
            pie_pagina(canvas, _doc)

            # Línea decorativa separadora (debajo del pie institucional)
            canvas.setStrokeColor(colors.HexColor("#BDC3C7"))
            canvas.setLineWidth(0.5)
            canvas.line(70, 25, page_width - 70, 25)
            
            # Metadata del sistema (debajo de la línea)
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(colors.HexColor("#7F8C8D"))
            
            # Obtener usuario actual si está disponible
            usuario_texto = "Usuario: "
            if parent and hasattr(parent, 'usuario_actual'):
                usuario_texto += parent.usuario_actual.get('username', 'Sistema')
            else:
                usuario_texto += "Sistema"
            
            # Izquierda: Generado por SIRA
            canvas.drawString(70, 17, "Generado por SIRA v1.0")
            
            # Centro: Usuario
            text_width = canvas.stringWidth(usuario_texto, "Helvetica", 7)
            canvas.drawString((page_width - text_width) / 2, 17, usuario_texto)
            
            # Derecha: Número de página
            canvas.drawRightString(page_width - 70, 17, f"Página {doc.page}")
            
            canvas.restoreState()

        # Construir PDF con footer personalizado
        doc.build(story, onFirstPage=footer_metadata, onLaterPages=footer_metadata)
        buf.close()
        
        return ruta
        
    except Exception as e:
        if parent:
            crear_msgbox(
                parent,
                "Error",
                f"No se pudo exportar el reporte: {e}",
                QMessageBox.Icon.Critical
            ).exec()
        return None


def exportar_tabla_excel(nombre_archivo: str, encabezados: list, filas: list) -> str:
    """Exporta datos tabulares a un archivo Excel."""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Datos"

        # Escribir encabezados
        ws.append(encabezados)

        # Escribir filas (convertir objetos date/datetime a strings)
        for fila in filas:
            fila_procesada = []
            for celda in fila:
                if isinstance(celda, (date, datetime)):
                    fila_procesada.append(celda.strftime("%d-%m-%Y"))
                elif celda is None:
                    fila_procesada.append("")
                else:
                    fila_procesada.append(str(celda))
            ws.append(fila_procesada)

        # Guardar archivo
        wb.save(nombre_archivo)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando archivo Excel: {e}")


def exportar_estudiantes_excel(parent, estudiantes: list) -> str:
    """Exporta lista de estudiantes a Excel."""
    if not estudiantes:
        if parent:
            crear_msgbox(
                parent,
                "Sin datos",
                "No hay estudiantes para exportar.",
                QMessageBox.Icon.Warning
            ).exec()
        return None

    try:
        sugerido = f"estudiantes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        ruta, _ = QFileDialog.getSaveFileName(
            parent,
            "Guardar reporte de estudiantes",
            sugerido,
            "Archivos Excel (*.xlsx)"
        )
        
        if not ruta:
            return None
            
        if not ruta.endswith(".xlsx"):
            ruta += ".xlsx"

        encabezados = list(estudiantes[0].keys())
        filas = [list(e.values()) for e in estudiantes]

        return exportar_tabla_excel(ruta, encabezados, filas)
        
    except Exception as e:
        if parent:
            crear_msgbox(
                parent,
                "Error",
                f"No se pudo exportar a Excel: {e}",
                QMessageBox.Icon.Critical
            ).exec()
        return None


def exportar_empleados_excel(parent, empleados: list) -> str:
    """Exporta lista de empleados a Excel."""
    if not empleados:
        if parent:
            crear_msgbox(
                parent,
                "Sin datos",
                "No hay empleados para exportar.",
                QMessageBox.Icon.Warning
            ).exec()
        return None

    try:
        sugerido = f"empleados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        ruta, _ = QFileDialog.getSaveFileName(
            parent,
            "Guardar reporte de empleados",
            sugerido,
            "Archivos Excel (*.xlsx)"
        )
        
        if not ruta:
            return None
            
        if not ruta.endswith(".xlsx"):
            ruta += ".xlsx"

        encabezados = list(empleados[0].keys())
        filas = [list(e.values()) for e in empleados]

        return exportar_tabla_excel(ruta, encabezados, filas)
        
    except Exception as e:
        if parent:
            crear_msgbox(
                parent,
                "Error",
                f"No se pudo exportar a Excel: {e}",
                QMessageBox.Icon.Critical
            ).exec()
        return None


def generar_certificado_promocion_sexto(estudiante: dict, institucion: dict, anio_escolar_egreso: str) -> str:
    """Genera certificado de promoción de 6to grado a 1er año de secundaria."""
    # Mapear campos de BD a formato esperado
    estudiante_norm = {"Nombres": estudiante.get("Nombres") or estudiante.get("nombres", ""),
                       "Apellidos": estudiante.get("Apellidos") or estudiante.get("apellidos", ""),
                       "Cédula": estudiante.get("Cédula") or estudiante.get("cedula", ""),
                       "Ciudad": estudiante.get("Ciudad") or estudiante.get("ciudad", ""),
                       "Fecha Nac.": estudiante.get("Fecha Nac.") or estudiante.get("fecha_nac")}

    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Ciudad", "Fecha Nac."]
    valido, mensaje = validar_datos_exportacion(estudiante_norm, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    if not anio_escolar_egreso:
        raise ValueError("Año escolar de egreso no proporcionado")
    
    # Normalizar datos
    estudiante_norm["Nombres"] = str(estudiante_norm["Nombres"]).strip().upper()
    estudiante_norm["Apellidos"] = str(estudiante_norm["Apellidos"]).strip().upper()
    cedula_normalizada = normalizar_cedula(estudiante_norm["Cédula"], es_estudiante=True)
    
    # Convertir fecha
    fecha_nac_str = convertir_fecha_string(estudiante_norm['Fecha Nac.'])
    
    # Extraer año escolar (formato: "2023/2024" -> "2023-2024")
    anio_periodo = str(anio_escolar_egreso).replace('/', '-')
    anios = anio_periodo.split('-')
    if len(anios) == 2:
        try:
            anio_inicio_egreso = int(anios[0].strip())
            anio_fin_egreso = int(anios[1].strip())
            anio_periodo_form = f"{formatear_anio(anio_inicio_egreso)}-{formatear_anio(anio_fin_egreso)}"
        except ValueError:
            anio_periodo_form = anio_periodo  # Mantener original si no se puede parsear
    else:
        anio_periodo_form = anio_periodo
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Certificados de promocion")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Certificado_prosecusion_primaria{estudiante_norm['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=180,
            bottomMargin=50
        )

        story = [Paragraph("CERTIFICADO DE EDUCACIÓN PRIMARIA", styles["Title"]), Spacer(1, 16)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        # Intentar obtener ultima_seccion desde ambos formatos
        ultima_seccion = (estudiante.get('ultima_seccion') or 
                         estudiante.get('Ultima Seccion') or 
                         estudiante.get('letra') or 'N/A')
        
        texto = (
            f"Quien suscribe <b>{institucion['director'].upper()}</b> titular de la Cédula de Identidad "
            f"Nº <b>{director_ci}</b> Director(a) de la Institución Educativa "
            f"<b>{institucion['nombre'].upper()}</b>, ubicada en el Municipio <b>JUAN ANTONIO SOTILLO</b> de la Parroquia "
            f"<b>PUERTO LA CRUZ</b> adscrita al Centro de Desarrollo de la Calidad Educativa Estadal <b>ANZOÁTEGUI</b>. "
            f"Por la presente certifica que el (la) estudiante <b>{estudiante_norm['Apellidos']} {estudiante_norm['Nombres']}</b> "
            f"titular de Cédula Escolar Nº <b>{cedula_normalizada}</b>, nacido (a) en el Municipio "
            f"<b>{estudiante_norm['Ciudad']}</b> del Estado <b>ANZOÁTEGUI</b>, en fecha <b>{fecha_nac_str}</b>, cursó el "
            f"<b>6to Grado</b> correspondiéndole el literal <b>{ultima_seccion}</b> "
            f"durante el periodo escolar <b>{anio_periodo_form}</b>, siendo promovido(a) al <b>1er Año del Nivel de "
            f"Educación Media</b>, previo cumplimiento a los requisitos establecidos en la normativa legal vigente."
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 40))

        # Fecha de expedición
        fecha_hoy = date.today()
        anio_actual = fecha_hoy.year
        anio_actual_form = formatear_anio(anio_actual)
        
        texto_fecha = f"Certificado que se expide en <b>PUERTO LA CRUZ</b>, a los <b>30</b> días del mes de <b>Julio</b> de <b>{anio_actual_form}</b>"
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 30))

        # Tabla de firmas institucionales
        tabla_firmas_data = [
            # Encabezados
            [
                Paragraph("<b>INSTITUCIÓN EDUCATIVA<br/>(PARA VALIDEZ NACIONAL)</b>", centrado),
                Paragraph("<b>CENTRO DE DESARROLLO DE LA CALIDAD<br/>EDUCATIVA ESTADAL<br/>(PARA VALIDEZ INTERNACIONAL)</b>", centrado)
            ],
            # Director(a)
            [
                Paragraph("DIRECTOR(A)", styles["Normal"]),
                Paragraph("DIRECTOR(A)", styles["Normal"])
            ],
            # Nombre y Apellido
            [
                Paragraph(f"Nombre y Apellido: {institucion['director'].upper()}", styles["Normal"]),
                Paragraph("Nombre y Apellido:", styles["Normal"])
            ],
            # Cédula de Identidad
            [
                Paragraph(f"Número de C.I: {director_ci}", styles["Normal"]),
                Paragraph("Número de C.I:", styles["Normal"])
            ],
            # Firma y Sello (espacios vacíos)
            [
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"]),
                Paragraph("Firma y Sello:<br/><br/><br/>", styles["Normal"])
            ]
        ]

        tabla_firmas = Table(tabla_firmas_data, colWidths=[page_width/2 - 100, page_width/2 - 100])
        tabla_firmas.setStyle(TableStyle([
            # Bordes externos de la tabla
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            
            # Alineación
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Encabezados centrados
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),   # Resto alineado a la izquierda
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),   # Alineación vertical superior
            
            # Padding interno
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(tabla_firmas)

        # Construir PDF con encabezado personalizado (sin pie de página)
        doc.build(story, onFirstPage=encabezado_prosecucion, onLaterPages=encabezado_prosecucion)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")

def generar_constancia_retiro(estudiante: dict, institucion: dict, anio_escolar: dict, motivo_retiro: str = None) -> str:
    """Genera constancia de retiro en PDF para un estudiante."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula", "Grado", "Ciudad", "Fecha Nac."]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(f"Datos de estudiante incompletos: {mensaje}")
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        raise ValueError(f"Datos de institución incompletos: {mensaje}")
    
    # Extraer años
    anio_inicio, anio_fin = extraer_anio_escolar(anio_escolar)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    estudiante["Cédula"] = normalizar_cedula(estudiante["Cédula"], es_estudiante=True)

    # Convertir fecha
    fecha_nac_str = convertir_fecha_string(estudiante['Fecha Nac.'])
    
    # Calcular edad
    try:
        if isinstance(estudiante['Fecha Nac.'], (date, datetime)):
            edad = calcular_edad(estudiante['Fecha Nac.'])
        else:
            fecha_obj = datetime.strptime(str(estudiante['Fecha Nac.']), "%d-%m-%Y").date()
            edad = calcular_edad(fecha_obj)
    except (ValueError, TypeError, KeyError):
        edad = "N/A"
    
    # Motivo por defecto si no se proporciona
    if not motivo_retiro:
        motivo_retiro = "es retirado de la institución a solicitud de su representante siendo Promovido al siguiente grado"
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Constancias de retiro")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)

    nombre_base = sanitizar_nombre_archivo(f"Constancia_retiro_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")

    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=220,
            bottomMargin=50
        )

        story = [Paragraph("CONSTANCIA DE RETIRO", styles["Title"]), Spacer(1, 20)]

        # Texto principal
        director_ci = normalizar_cedula(institucion['director_ci'])
        
        # Determinar género para el artículo
        genero = estudiante.get("Género", "").lower()
        articulo = "el" if genero == "masculino" else "la"
        
        # Determinar si es nivel (Inicial) o grado (Primaria)
        tipo_educacion = estudiante.get("Tipo Educ.", "").lower()
        if "inicial" in tipo_educacion:
            grado_texto = f"{estudiante['Grado']}"
        else:
            grado_texto = f"{estudiante['Grado']} grado"
            
        texto = (
            f"La Dirección del plantel hace constar mediante la presente, que {articulo} Alumno(a): "
            f"<b>{estudiante['Apellidos']} {estudiante['Nombres']}</b>, "
            f"nacido(a) en <b>{estudiante['Ciudad'].upper()}</b> él <b>{fecha_nac_str}</b> "
            f"y de <b>{edad} año(s)</b> de edad, estudiante regular del <b>{grado_texto}</b> "
            f"para el año escolar <b>{anio_inicio}-{anio_fin}</b>, {motivo_retiro}.<br/><br/>"
        )
        story.append(Paragraph(texto, justificado))
        story.append(Spacer(1, 20))

        # Fecha actual
        fecha_hoy = date.today()
        dia = fecha_hoy.day
        mes_nombre = fecha_hoy.strftime("%B").upper()
        
        # Traducir mes al español
        meses = {
            'JANUARY': 'ENERO', 'FEBRUARY': 'FEBRERO', 'MARCH': 'MARZO',
            'APRIL': 'ABRIL', 'MAY': 'MAYO', 'JUNE': 'JUNIO',
            'JULY': 'JULIO', 'AUGUST': 'AGOSTO', 'SEPTEMBER': 'SEPTIEMBRE',
            'OCTOBER': 'OCTUBRE', 'NOVEMBER': 'NOVIEMBRE', 'DECEMBER': 'DICIEMBRE'
        }
        mes_es = meses.get(mes_nombre, mes_nombre)
        anio = fecha_hoy.year
        
        texto_fecha = (
            f"Se expide constancia a solicitud de la parte interesada en Puerto La Cruz "
            f"a los <b>{dia}</b> días del mes de <b>{mes_es}</b> del año <b>{anio}</b>."
        )
        story.append(Paragraph(texto_fecha, justificado))
        story.append(Spacer(1, 150))

        # Firma
        firma_texto = f"________________________<br/>Prof. {institucion['director']}"
        story.append(Paragraph(firma_texto, centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph(f"C.I. {director_ci}", centrado))
        story.append(Spacer(1, 3))
        story.append(Paragraph("Director", centrado))

        # Construir PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando PDF: {e}")


def generar_historial_estudiante_pdf(estudiante: dict, historial: list) -> str:
    """Genera un PDF con el historial académico completo del estudiante."""
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(mensaje)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    
    # Normalizar cédula escolar (agregar CE- si no lo tiene)
    cedula = str(estudiante["Cédula"]).strip()
    if not cedula.upper().startswith("CE-"):
        cedula_escolar = f"CE-{cedula}"
    else:
        cedula_escolar = cedula.upper()
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Historial academico")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)
    
    nombre_base = sanitizar_nombre_archivo(f"Historial_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")
    
    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=180,
            bottomMargin=50
        )

        story = [Paragraph("HISTORIAL ACADÉMICO", styles["Title"]), Spacer(1, 16)]

        # Datos del estudiante
        datos_estudiante = f"""
        <b>Estudiante:</b> {estudiante['Nombres']} {estudiante['Apellidos']}<br/>
        <b>Cédula Escolar:</b> {cedula_escolar}<br/>
        """
        
        # Agregar grado actual si existe
        if "Grado" in estudiante and estudiante["Grado"]:
            datos_estudiante += f"<b>Grado actual:</b> {estudiante['Grado']}<br/>"
        
        story.append(Paragraph(datos_estudiante, styles["Normal"]))
        story.append(Spacer(1, 20))
        
        # Tabla de historial
        if historial:
            # Encabezados de la tabla
            datos_historial = [
                [
                    Paragraph("<b>Año Escolar</b>", centrado),
                    Paragraph("<b>Nivel</b>", centrado),
                    Paragraph("<b>Grado</b>", centrado),
                    Paragraph("<b>Sección</b>", centrado),
                    Paragraph("<b>Docente</b>", centrado)
                ]
            ]
            
            # Agregar filas del historial
            for registro in historial:
                anio_escolar = f"{registro['año_inicio']}-{registro['año_inicio']+1}"
                datos_historial.append([
                    Paragraph(anio_escolar, centrado),
                    Paragraph(str(registro['nivel']), centrado),
                    Paragraph(str(registro['grado']), centrado),
                    Paragraph(str(registro['letra']), centrado),
                    Paragraph(str(registro.get('docente', 'Sin asignar')), centrado)
                ])
            
            tabla_historial = Table(datos_historial, colWidths=[90, 110, 70, 70, 130])
            tabla_historial.setStyle(TableStyle([
                # Encabezado
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E5894')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                
                # Contenido
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                
                # Bordes y relleno
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            story.append(tabla_historial)
        else:
            # Sin historial
            texto_sin_datos = Paragraph(
                "<i>No se encontró historial académico para este estudiante.</i>",
                justificado
            )
            story.append(texto_sin_datos)
        
        # Generar PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando historial en PDF: {e}")


def generar_listado_estudiantes_seccion(seccion: dict, estudiantes: list, docente: dict, institucion: dict) -> str:
    """Genera un listado en PDF de los estudiantes activos de una sección."""
    # Validar datos
    if not seccion:
        return "Error: No se proporcionaron datos de la sección."
    
    if not estudiantes:
        return "Error: No hay estudiantes para exportar."
    
    campos_inst = ["director", "director_ci", "nombre"]
    valido, mensaje = validar_datos_exportacion(institucion, campos_inst)
    if not valido:
        return mensaje
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Listados de secciones")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        return msg
    
    # Nombre de archivo
    nivel = seccion.get('nivel', 'Nivel')
    grado = seccion.get('grado', 'Grado')
    letra = seccion.get('letra', 'Seccion')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_base = sanitizar_nombre_archivo(f"Listado_{nivel}_{grado}_{letra}_{timestamp}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")
    
    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=80,
            rightMargin=80,
            topMargin=180,
            bottomMargin=50
        )
        
        elementos = []
        
        # Título
        titulo_style = ParagraphStyle(
            name="TituloListado",
            parent=styles["Heading1"],
            alignment=TA_CENTER,
            fontSize=12,
            spaceAfter=8,
            textColor=colors.HexColor("#2c3e50")
        )
        
        titulo_texto = f"LISTADO DE ESTUDIANTES<br/>{nivel} - {grado} Sección {letra}"
        titulo = Paragraph(titulo_texto, titulo_style)
        elementos.append(titulo)
        elementos.append(Spacer(1, 4))
        
        if docente and docente.get('nombre_completo'):
            docente_nombre = str(docente['nombre_completo']).strip().upper()
        else:
            docente_nombre = "SIN DOCENTE ASIGNADO"
        
        docente_style_izq = ParagraphStyle(
            name="DocenteIzq",
            parent=styles["Normal"],
            fontSize=9,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#333333")
        )
        
        total_style_der = ParagraphStyle(
            name="TotalDer",
            parent=styles["Normal"],
            fontSize=9,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#333333")
        )
        
        docente_texto = f"<b>Docente:</b> {docente_nombre}"
        total_texto = f"<b>Total de estudiantes: {len(estudiantes)}</b>"
        
        # Tabla sin bordes para alinear docente (izq) y total (der)
        info_tabla = Table(
            [[Paragraph(docente_texto, docente_style_izq), Paragraph(total_texto, total_style_der)]],
            colWidths=[300, 185]
        )
        info_tabla.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        elementos.append(info_tabla)
        elementos.append(Spacer(1, 8))
        
        # Crear tabla de estudiantes
        # Encabezados
        datos_tabla = [[
            "N°",
            "Cédula Escolar",
            "Apellidos",
            "Nombres",
            "Edad",
            "Género",
            "Observaciones"
        ]]
        
        # Agregar estudiantes
        for idx, est in enumerate(estudiantes, start=1):
            cedula = est.get('cedula', 'N/A')
            apellidos = str(est.get('apellidos', '')).strip().upper()
            nombres = str(est.get('nombres', '')).strip().upper()
            
            edad = est.get('edad', None)
            if edad is not None and isinstance(edad, (int, float)):
                edad_str = f"{int(edad)} años"
            else:
                edad_str = "N/A"
            
            genero = est.get('genero', 'N/A')
            if isinstance(genero, str):
                genero = genero.upper()
            
            datos_tabla.append([
                str(idx),
                str(cedula),
                apellidos,
                nombres,
                edad_str,
                str(genero)
            ])
        
        # Crear tabla
        tabla = Table(datos_tabla, colWidths=[25, 90, 130, 130, 55, 40, 80], rowHeights=13)
        
        tabla.setStyle(TableStyle([
            # Encabezado
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1a5490")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
            
            # Datos
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 7.5),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Número centrado
            ('ALIGN', (1, 1), (1, -1), 'CENTER'),  # Cédula centrada
            ('ALIGN', (4, 1), (4, -1), 'CENTER'),  # Edad centrada
            ('ALIGN', (5, 1), (5, -1), 'CENTER'),  # Género centrado
            ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
            
            # Bordes
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor("#1a5490")),
            
            # Alternancia de colores en filas
            *[('BACKGROUND', (0, i), (-1, i), colors.HexColor("#f0f0f0")) 
              for i in range(2, len(datos_tabla), 2)]
        ]))
        
        elementos.append(tabla)
        
        # Construir PDF
        doc.build(elementos, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        
        return nombre_archivo
        
    except Exception as e:
        return f"Error generando listado: {e}"


def generar_reporte_rac(parent, empleados: list, institucion: dict) -> str:
    """Genera reporte RAC (Registro de Asignación de Cargos) en formato Excel."""
    if not empleados:
        crear_msgbox(
            parent,
            "Sin datos",
            "No hay empleados para exportar",
            QMessageBox.Icon.Warning
        ).exec()
        return None
    
    try:
        # Diálogo para guardar archivo
        fecha_actual = datetime.now().strftime("%Y%m%d")
        nombre_sugerido = f"Reporte_RAC_{fecha_actual}.xlsx"
        
        archivo, _ = QFileDialog.getSaveFileName(
            parent,
            "Guardar Reporte RAC",
            nombre_sugerido,
            "Archivos Excel (*.xlsx)"
        )
        
        if not archivo:
            return None
        
        # Crear workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "RAC"
        
        # Contar secciones activas del año escolar actual para especialistas
        total_secciones_activas = SeccionesModel.contar_activas_anio_actual()
        
        # Datos fijos de la institución (valores predeterminados)
        datos_fijos = {
            'cod_estado': '2',
            'estado': 'ANZOATEGUI',
            'municipio': 'JUAN ANTONIO SOTILLO',
            'parroquia': 'PUERTO LA CRUZ',
            'codigo_dependencia': institucion.get('codigo_dependencia'),
            'codigo_estadistico': institucion.get('codigo_estadistico'),
            'codigo_plantel': institucion.get('codigo_dea'),
            'nombre_plantel': institucion.get('nombre').upper(),
            'nivel': 'PRIMARIA',
            'modalidad': 'PRIMARIA',
            'ubicacion': 'URBANA',
            'turnos_plantel': 'INTEGRAL'
        }
        
        # Encabezados
        encabezados = [
            'COD ESTADO', 'ESTADO', 'MUNICIPIO', 'PARROQUIA',
            'CODIGO DEPENDENCIA', 'CODIGO ESTADISTICO', 'CODIGO DEL PLANTEL',
            'NOMBRE DEL PLANTEL EN NOMINA', 'NIVEL', 'MODALIDAD',
            'UBICACIÓN GEOGRAFICA', 'TURNOS QUE ATIENDE EL PLANTEL',
            'CODIGO RAC', 'CARGO', 'TIPO PERSONAL', 'CEDULA',
            'NOMBRE Y APELLIDO', 'FECHA DE NACIMIENTO', 'LUGAR DE NACIMIENTO',
            'TELÉFONO', 'CORREO', 'NIVEL DE INSTRUCCION', 'PROFESIÓN','TALLA CAMISA',
            'TALLA PANTALÓN', 'TALLA ZAPATOS', 
            '¿PRACTICA ALGUNA ACTIVIDAD REGULARMENTE?. DE SER POSITIVO COLOCAR CUÁL',
            '¿PRACTICA CULTURAL REGULARMENTE?', 'TIPO DE VIVIENDA (CASA, APARTAMENTO, QUINTA, RANCHO)',
            'CONDICIÓN DE LA VIVIENDA (PROPIA, ALQUILADA, DE UN FAMILIAR U OTRO)',
            'TIPO DE MATERIAL (EN CASO DE NECESITAR REPARACIONES, ESPECIFIQUE EL TIPO DE MATERIAL REQUERIDO)',
            'TIPO DE ENFERMEDAD, ¿PADECE DE ALGUNA ENFERMEDAD? DE SER POSITIVO INDIQUE CUÁL',
            'MEDICAMENTOS, ¿REQUIERE ALGÚN MEDICAMENTO DE USO?',
            '¿POSEE ALGUNA DISCAPACIDAD CERTIFICADA? DE SER POSITIVO INDIQUE CUAL',
            'FECHA DE INGRESO', 'SEXO',
            'HORAS ACADEMICAS', 'HORAS ADM', 'TURNO QUE ATIENDE',
            'GRADO QUE IMPARTE EL DOCENTE', 'SECCIÓN',
            'ESPECIALIDAD QUE IMPARTE EL DOCENTE', 'AÑO', 'SECCIONES',
            'MATERIA QUE IMPARTE O ESPECIALIDAD', 'PERIODO O GRUPO',
            'SITUACIÓN DEL TRABAJADOR', 'OBSERVACIÓN'
        ]
        
        ws.append(encabezados)
        
        # Función auxiliar para formatear fecha o dejar vacío
        def _fmt_fecha(valor):
            if valor is None:
                return ''
            if isinstance(valor, (date, datetime)):
                return valor.strftime('%d-%m-%Y')
            return str(valor).strip() if valor else ''
        
        # Función auxiliar: devolver valor de BD como string o vacío
        def _val(_empleado, campo):
            v = _empleado.get(campo)
            if v is None:
                return ''
            v = str(v).strip()
            return v if v else ''
        
        # Procesar cada empleado
        for empleado in empleados:
            # Normalizar cédula (quitar prefijo V-/E- y puntos)
            cedula = _val(empleado, 'cedula')
            if cedula.startswith(('V-', 'E-', 'J-', 'G-')):
                cedula = cedula[2:]
            cedula = cedula.replace('.', '').replace('-', '')
            
            # Nombre completo
            nombres = _val(empleado, 'nombres').upper()
            apellidos = _val(empleado, 'apellidos').upper()
            nombre_completo = f"{nombres} {apellidos}".strip()
            
            # Cargo y tipo de personal
            cargo = _val(empleado, 'cargo').upper()
            tipo_personal = _val(empleado, 'tipo_personal')
            
            # Código RAC
            codigo_rac = _val(empleado, 'codigo_rac')
            
            # Fechas
            fecha_nac_str = _fmt_fecha(empleado.get('fecha_nac'))
            fecha_ingreso_str = _fmt_fecha(empleado.get('fecha_ingreso'))
            
            # Sexo
            sexo = _val(empleado, 'genero').upper()
            
            # Horas
            horas_acad_bd = empleado.get('horas_acad')
            horas_adm_bd = empleado.get('horas_adm')
            horas_acad = str(horas_acad_bd).replace('.', ',') if horas_acad_bd is not None else ''
            horas_adm = str(horas_adm_bd).replace('.', ',') if horas_adm_bd is not None else ''
            
            # Datos personales adicionales
            lugar_nacimiento = _val(empleado, 'lugar_nacimiento').upper()
            telefono = _val(empleado, 'num_contact')
            correo = _val(empleado, 'correo')
            nivel_instruccion = _val(empleado, 'nivel_instruccion').upper()
            profesion = _val(empleado, 'profesion').upper()
            talla_camisa = _val(empleado, 'talla_camisa')
            talla_pantalon = _val(empleado, 'talla_pantalon')
            talla_zapatos = _val(empleado, 'talla_zapatos')
            actividad = _val(empleado, 'actividad')
            cultural = _val(empleado, 'cultural')
            tipo_vivienda = _val(empleado, 'tipo_vivienda')
            condicion_vivienda = _val(empleado, 'condicion_vivienda')
            material_vivienda = _val(empleado, 'material_vivienda')
            tipo_enfermedad = _val(empleado, 'tipo_enfermedad')
            medicamento = _val(empleado, 'medicamento')
            discapacidad = _val(empleado, 'discapacidad')
            
            # Estado del trabajador
            estado = empleado.get('estado', '')
            if isinstance(estado, int):
                situacion = 'ACTIVO' if estado == 1 else 'INACTIVO'
            elif estado:
                situacion = str(estado).upper()
            else:
                situacion = ''
            
            # Turno, grado y sección
            turno = 'INTEGRAL'
            grado_imparte = ''
            seccion_imparte = ''
            especialidad_imparte = ''
            
            especialidad = empleado.get('especialidad')
            
            if especialidad and str(especialidad).strip():
                especialidad_imparte = str(especialidad).strip().upper()
                grado_imparte = f"ESPECIALISTA EN {especialidad_imparte}"
                if total_secciones_activas > 0:
                    seccion_imparte = f"{total_secciones_activas} secciones"
                else:
                    seccion_imparte = "Todas"
            else:
                seccion_grado = empleado.get('seccion_grado')
                seccion_letra = empleado.get('seccion_letra')
                seccion_nivel = empleado.get('seccion_nivel')
                
                if seccion_grado and seccion_letra and seccion_nivel:
                    grado_imparte = seccion_grado
                    seccion_imparte = 'U' if str(seccion_letra).upper() == 'ÚNICA' else seccion_letra
                elif 'DIRECTOR' in cargo:
                    grado_imparte = 'DIRECTOR'
            
            # Construir fila alineada con los 48 encabezados
            fila = [
                datos_fijos['cod_estado'],           # 1  COD ESTADO
                datos_fijos['estado'],                # 2  ESTADO
                datos_fijos['municipio'],             # 3  MUNICIPIO
                datos_fijos['parroquia'],             # 4  PARROQUIA
                datos_fijos['codigo_dependencia'],    # 5  CODIGO DEPENDENCIA
                datos_fijos['codigo_estadistico'],    # 6  CODIGO ESTADISTICO
                datos_fijos['codigo_plantel'],        # 7  CODIGO DEL PLANTEL
                datos_fijos['nombre_plantel'],        # 8  NOMBRE DEL PLANTEL
                datos_fijos['nivel'],                 # 9  NIVEL
                datos_fijos['modalidad'],             # 10 MODALIDAD
                datos_fijos['ubicacion'],             # 11 UBICACIÓN GEOGRAFICA
                datos_fijos['turnos_plantel'],        # 12 TURNOS PLANTEL
                codigo_rac,                           # 13 CODIGO RAC
                cargo,                                # 14 CARGO
                tipo_personal,                        # 15 TIPO PERSONAL
                cedula,                               # 16 CEDULA
                nombre_completo,                      # 17 NOMBRE Y APELLIDO
                fecha_nac_str,                        # 18 FECHA DE NACIMIENTO
                lugar_nacimiento,                     # 19 LUGAR DE NACIMIENTO
                telefono,                             # 20 TELÉFONO
                correo,                               # 21 CORREO
                nivel_instruccion,                    # 22 NIVEL DE INSTRUCCION
                profesion,                            # 23 PROFESIÓN
                talla_camisa,                         # 24 TALLA CAMISA
                talla_pantalon,                       # 25 TALLA PANTALÓN
                talla_zapatos,                        # 26 TALLA ZAPATOS
                actividad,                            # 27 ACTIVIDAD
                cultural,                             # 28 CULTURAL
                tipo_vivienda,                        # 29 TIPO DE VIVIENDA
                condicion_vivienda,                   # 30 CONDICIÓN DE LA VIVIENDA
                material_vivienda,                    # 31 TIPO DE MATERIAL
                tipo_enfermedad,                      # 32 TIPO DE ENFERMEDAD
                medicamento,                          # 33 MEDICAMENTOS
                discapacidad,                         # 34 DISCAPACIDAD
                fecha_ingreso_str,                    # 35 FECHA DE INGRESO
                sexo,                                 # 36 SEXO
                horas_acad,                           # 37 HORAS ACADEMICAS
                horas_adm,                            # 38 HORAS ADM
                turno,                                # 39 TURNO QUE ATIENDE
                grado_imparte,                        # 40 GRADO QUE IMPARTE
                seccion_imparte,                      # 41 SECCIÓN
                especialidad_imparte,                 # 42 ESPECIALIDAD QUE IMPARTE
                '',                                   # 43 AÑO
                '',                                   # 44 SECCIONES
                '',                                   # 45 MATERIA
                '',                                   # 46 PERIODO O GRUPO
                situacion,                            # 47 SITUACIÓN DEL TRABAJADOR
                ''                                    # 48 OBSERVACIÓN
            ]
            
            ws.append(fila)
        
        # Aplicar estilos al reporte RAC
        # Colores y estilos
        azul_fondo_encabezado = PatternFill(start_color="4BACC6", end_color="4BACC6", fill_type="solid")
        azul_fondo_filas = PatternFill(start_color="B9CDE5", end_color="B9CDE5", fill_type="solid")
        blanco_texto = Font(name='Arial', size=12, color="FFFFFF", bold=True)
        negro_texto = Font(name='Arial', size=12, color="000000")
        borde_delgado = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style='thin', color='000000'),
            bottom=Side(style='thin', color='000000')
        )
        alineacion_centro = Alignment(horizontal='center', vertical='center', wrap_text=True)
        alineacion_izquierda = Alignment(horizontal='left', vertical='center', wrap_text=True)
        
        # Anchos de columna personalizados (en caracteres) — 48 columnas
        anchos_columnas = [
            10,  # 1  COD ESTADO
            15,  # 2  ESTADO
            20,  # 3  MUNICIPIO
            20,  # 4  PARROQUIA
            18,  # 5  CODIGO DEPENDENCIA
            18,  # 6  CODIGO ESTADISTICO
            18,  # 7  CODIGO DEL PLANTEL
            35,  # 8  NOMBRE DEL PLANTEL
            12,  # 9  NIVEL
            12,  # 10 MODALIDAD
            15,  # 11 UBICACIÓN
            15,  # 12 TURNOS PLANTEL
            12,  # 13 CODIGO RAC
            25,  # 14 CARGO
            12,  # 15 TIPO PERSONAL
            12,  # 16 CEDULA
            35,  # 17 NOMBRE Y APELLIDO
            15,  # 18 FECHA DE NACIMIENTO
            22,  # 19 LUGAR DE NACIMIENTO
            14,  # 20 TELÉFONO
            25,  # 21 CORREO
            22,  # 22 NIVEL DE INSTRUCCION
            20,  # 23 PROFESIÓN
            12,  # 24 TALLA CAMISA
            14,  # 25 TALLA PANTALÓN
            12,  # 26 TALLA ZAPATOS
            30,  # 27 ACTIVIDAD
            25,  # 28 CULTURAL
            20,  # 29 TIPO DE VIVIENDA
            25,  # 30 CONDICIÓN DE LA VIVIENDA
            30,  # 31 TIPO DE MATERIAL
            30,  # 32 TIPO DE ENFERMEDAD
            25,  # 33 MEDICAMENTOS
            30,  # 34 DISCAPACIDAD
            15,  # 35 FECHA DE INGRESO
            8,   # 36 SEXO
            15,  # 37 HORAS ACADEMICAS
            12,  # 38 HORAS ADM
            15,  # 39 TURNO QUE ATIENDE
            25,  # 40 GRADO QUE IMPARTE
            12,  # 41 SECCIÓN
            30,  # 42 ESPECIALIDAD
            8,   # 43 AÑO
            12,  # 44 SECCIONES
            30,  # 45 MATERIA
            15,  # 46 PERIODO
            15,  # 47 SITUACIÓN
            25   # 48 OBSERVACIÓN
        ]
        
        # Aplicar anchos de columna
        for idx, ancho in enumerate(anchos_columnas, start=1):
            letra_columna = ws.cell(row=1, column=idx).column_letter
            ws.column_dimensions[letra_columna].width = ancho
        
        # Estilizar encabezado (fila 1)
        for col in range(1, len(encabezados) + 1):
            celda = ws.cell(row=1, column=col)
            celda.fill = azul_fondo_encabezado
            celda.font = blanco_texto
            celda.alignment = alineacion_centro
            celda.border = borde_delgado
        
        # Estilizar filas de datos
        for fila_idx in range(2, ws.max_row + 1):
            ws.row_dimensions[fila_idx].height = 35  # Altura de fila
            for col in range(1, len(encabezados) + 1):
                celda = ws.cell(row=fila_idx, column=col)
                celda.fill = azul_fondo_filas
                celda.font = negro_texto
                celda.border = borde_delgado
                
                # Alineación según columna
                if col in [16, 17]:  # CEDULA, NOMBRE Y APELLIDO
                    celda.alignment = alineacion_izquierda
                else:
                    celda.alignment = alineacion_centro
        
        # Altura de la fila de encabezado
        ws.row_dimensions[1].height = 85
        
        # Guardar archivo
        wb.save(archivo)
        
        crear_msgbox(
            parent,
            "Exportación exitosa",
            f"Reporte RAC generado correctamente.\n\nTotal de empleados: {len(empleados)}\nArchivo: {os.path.basename(archivo)}",
            QMessageBox.Icon.Information
        ).exec()
        
        return archivo
        
    except Exception as e:
        crear_msgbox(
            parent,
            "Error",
            f"Error al generar reporte RAC:\n{str(e)}",
            QMessageBox.Icon.Critical
        ).exec()
        return None

def generar_cuadratura_excel(parent) -> str:
    """Genera la Cuadratura Maternal, Inicial y Primaria en formato Excel"""
    try:
        anio = AnioEscolarModel.obtener_actual()
        if not anio:
            crear_msgbox(
                parent, "Error",
                "No se pudo obtener el año escolar actual.",
                QMessageBox.Icon.Warning
            ).exec()
            return None

        secciones = SeccionesModel.obtener_todas(anio['id'], solo_activas=True)
        empleados = EmpleadoModel.listar_activos()
        institucion = InstitucionModel.obtener_por_id(1)

        if not empleados:
            crear_msgbox(
                parent, "Sin datos",
                "No hay empleados activos para generar la cuadratura.",
                QMessageBox.Icon.Warning
            ).exec()
            return None

        fecha_actual = datetime.now().strftime("%Y%m%d")
        nombre_sugerido = f"Cuadratura_{fecha_actual}.xlsx"
        archivo, _ = QFileDialog.getSaveFileName(
            parent, "Guardar Cuadratura",
            nombre_sugerido, "Archivos Excel (*.xlsx)"
        )
        if not archivo:
            return None

        codigo_dea = institucion.get('codigo_dea', '') if institucion else ''
        nombre_plantel = institucion.get('nombre', '') if institucion else ''

        GRADO_MAP = {
            'MATERNAL': {
                'label': 'MATERNAL', 'niveles': [], 'grados': [],
                'section_prefix': 'GRUPO', 'max_sections': 4,
            },
            'INICIAL': {
                'label': 'INICIAL', 'niveles': ['Inicial'],
                'grados': SeccionesModel.GRADOS_INICIAL,
                'section_prefix': 'SECCION', 'max_sections': 5,
            },
            '1ERO': {
                'label': '1 ER GRADO', 'niveles': ['Primaria'],
                'grados': ['1ero'], 'section_prefix': 'SECCION',
                'max_sections': 5,
            },
            '2DO': {
                'label': '2DO GRADO', 'niveles': ['Primaria'],
                'grados': ['2do'], 'section_prefix': 'SECCION',
                'max_sections': 5,
            },
            '3ERO': {
                'label': '3 ER GRADO', 'niveles': ['Primaria'],
                'grados': ['3ero'], 'section_prefix': 'SECCION',
                'max_sections': 5,
            },
            '4TO': {
                'label': '4TO GRADO', 'niveles': ['Primaria'],
                'grados': ['4to'], 'section_prefix': 'SECCION',
                'max_sections': 5,
            },
            '5TO': {
                'label': '5TO GRADO', 'niveles': ['Primaria'],
                'grados': ['5to'], 'section_prefix': 'SECCION',
                'max_sections': 5,
            },
        '6TO': {
            'label': '6TO GRADO', 'niveles': ['Primaria'],
            'grados': ['6to'], 'section_prefix': 'SECCION',
            'max_sections': 5,
        },
    }
        BLOCK_ORDER = ['MATERNAL', 'INICIAL', '1ERO', '2DO', '3ERO', '4TO', '5TO', '6TO']

        def _secciones_del_bloque(bloque_key):
            """Obtiene las secciones de BD que corresponden a un bloque."""
            cfg = GRADO_MAP[bloque_key]
            if bloque_key == 'MATERNAL':
                return []
            out = []
            for s in secciones:
                if s['nivel'] in cfg['niveles'] and s['grado'] in cfg['grados']:
                    out.append(s)
            return out

        def _letras_secciones_bloque(bloque_key):
            """Devuelve lista de letras de sección para un bloque (ej: ['A','B'])."""
            s_bd = _secciones_del_bloque(bloque_key)
            if not s_bd:
                return []
            letras = sorted(set(s['letra'] for s in s_bd))
            return letras

        bloque_secciones = {}
        bloque_letras = {}
        bloque_col_start = {}
        bloque_col_end = {}
        bloque_grade_cols = {}
        bloque_sec_cols = {}

        col_cursor = 8
        for bk in BLOCK_ORDER:
            cfg = GRADO_MAP[bk]
            letras = _letras_secciones_bloque(bk)
            n_grade = min(max(len(letras), 1), cfg['max_sections'])
            n_sec = min(max(len(letras), 1), cfg['max_sections'])

            bloque_secciones[bk] = _secciones_del_bloque(bk)
            bloque_letras[bk] = letras
            bloque_col_start[bk] = col_cursor
            bloque_grade_cols[bk] = (col_cursor, col_cursor + n_grade - 1)

            if bk == 'MATERNAL':
                bloque_sec_cols[bk] = None
                bloque_col_end[bk] = col_cursor + n_grade - 1
                col_cursor = bloque_col_end[bk] + 1
            else:
                sec_start = col_cursor + n_grade + 1
                bloque_sec_cols[bk] = (sec_start, sec_start + n_sec - 1)
                bloque_col_end[bk] = sec_start + n_sec - 1
                col_cursor = bloque_col_end[bk] + 1

        last_data_col = bloque_col_end[BLOCK_ORDER[-1]]

        wb = Workbook()
        ws = wb.active
        ws.title = "Cuadratura"

        font_title = Font(name='Arial', size=12, bold=True)
        font_label = Font(name='Arial', size=12, bold=True)
        font_section_row5 = Font(name='Arial', size=12, bold=True)
        font_section_row6 = Font(name='Arial', size=10, bold=True)
        font_section_row6_small = Font(name='Arial', size=8, bold=True)
        font_row7 = Font(name='Arial', size=10, bold=True)
        font_data = Font(name='Arial', size=10)
        font_data_bold = Font(name='Arial', size=10, bold=True)

        align_center = Alignment(horizontal='center', vertical='center')
        align_center_rot90 = Alignment(horizontal='center', vertical='center', text_rotation=90)
        align_left_rot90 = Alignment(horizontal='left', vertical='center', text_rotation=90)
        align_left = Alignment(horizontal='left', vertical='center')
        align_center_wrap = Alignment(horizontal='center', vertical='center', wrap_text=True)

        side_thick = Side(style='thick')
        side_thin = Side(style='thin')
        side_medium = Side(style='medium')
        border_header_full = Border(
            left=side_thick, right=side_thin,
            top=side_thick, bottom=side_thick
        )
        border_header_right = Border(
            left=side_thin, right=side_thick,
            top=side_thick, bottom=side_thick
        )
        border_data_full = Border(
            left=side_thick, right=side_thin,
            top=side_thin, bottom=side_thin
        )
        border_data_right = Border(
            left=side_thin, right=side_thick,
            top=side_thin, bottom=side_thin
        )
        border_data_all_thin = Border(
            left=side_thin, right=side_thin,
            top=side_thin, bottom=side_thin
        )
        border_data_first_last = Border(
            left=side_thick, right=side_thick,
            top=side_thin, bottom=side_thin
        )
        border_matricula = Border(
            left=side_thin, right=side_thin,
            top=side_medium, bottom=side_medium
        )

        light_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")

        # ── Row 2: Título ──
        ws.cell(row=2, column=8, value="CUADRATURA MATERNAL, INICIAL Y ").font = font_title
        ws.cell(row=2, column=8).alignment = align_center
        ws.merge_cells(
            start_row=2, start_column=8,
            end_row=2, end_column=last_data_col
        )

        # ── Row 3: Info institucional ──
        col3 = 8
        ws.cell(row=3, column=col3, value="Codigo Interno del Plantel").font = font_label
        ws.cell(row=3, column=col3).alignment = align_center
        cod_end = min(col3 + 8, last_data_col)
        ws.merge_cells(start_row=3, start_column=col3, end_row=3, end_column=cod_end)

        col3 = cod_end + 1
        ws.cell(row=3, column=col3, value=str(codigo_dea)).font = font_label
        ws.cell(row=3, column=col3).alignment = align_center
        dea_end = min(col3 + 5, last_data_col)
        ws.merge_cells(start_row=3, start_column=col3, end_row=3, end_column=dea_end)

        col3 = dea_end + 1
        ws.cell(row=3, column=col3, value="Nombre del Plantel:").font = font_label
        ws.cell(row=3, column=col3).alignment = align_center
        np_end = min(col3 + 6, last_data_col)
        ws.merge_cells(start_row=3, start_column=col3, end_row=3, end_column=np_end)

        col3 = np_end + 1
        ws.cell(row=3, column=col3, value=nombre_plantel.upper()).font = font_label
        ws.cell(row=3, column=col3).alignment = Alignment(horizontal='left', vertical='center')
        if col3 <= last_data_col:
            ws.merge_cells(start_row=3, start_column=col3, end_row=3, end_column=last_data_col)

        # ── Row 4: Vacía ──

        # ── Row 5: Grade labels + SECCIONES labels ──
        # ── Row 6: Section sub-labels (rotated 90°) ──
        # ── Row 7: C.I. / Nombre del Docente / MATRICULA ──

        for bk in BLOCK_ORDER:
            cfg = GRADO_MAP[bk]
            letras = bloque_letras[bk]
            cs = bloque_col_start[bk]
            grade_start, grade_end = bloque_grade_cols[bk]

            n_grade = grade_end - grade_start + 1

            # Row 5: grade label
            ws.cell(row=5, column=grade_start, value=cfg['label']).font = font_section_row5
            ws.cell(row=5, column=grade_start).alignment = align_center
            if grade_end > grade_start:
                ws.merge_cells(
                    start_row=5, start_column=grade_start,
                    end_row=5, end_column=grade_end
                )

            # Row 6: section sub-labels
            for i in range(n_grade):
                col = grade_start + i
                if cfg['section_prefix'] == 'GRUPO':
                    label = f"GRUPO {i + 1}"
                else:
                    if i < len(letras):
                        label = f"SECCION {letras[i]}"
                    else:
                        label = f"SECCION {chr(65 + i)}"
                ws.cell(row=6, column=col, value=label).font = font_section_row6
                ws.cell(row=6, column=col).alignment = align_left_rot90

            # Row 7: MATRICULA (merged across grade + gap + secciones)
            if bloque_sec_cols[bk] is not None:
                sec_start, sec_end = bloque_sec_cols[bk]
                n_sec = sec_end - sec_start + 1

                # Row 5: SECCIONES label
                ws.cell(row=5, column=sec_start, value="SECCIONES").font = font_section_row5
                ws.cell(row=5, column=sec_start).alignment = align_center
                if sec_end > sec_start:
                    ws.merge_cells(
                        start_row=5, start_column=sec_start,
                        end_row=5, end_column=sec_end
                    )

                # Row 6: secciones sub-labels (empty, just formatting)
                for i in range(n_sec):
                    col = sec_start + i
                    ws.cell(row=6, column=col).font = font_section_row6
                    ws.cell(row=6, column=col).alignment = align_left_rot90

                # Row 7: MATRICULA merged across all cols in block
                mat_start = grade_start
                mat_end = sec_end
                ws.cell(row=7, column=mat_start, value="MATRICULA").font = font_row7
                ws.cell(row=7, column=mat_start).alignment = align_center
                if mat_end > mat_start:
                    ws.merge_cells(
                        start_row=7, start_column=mat_start,
                        end_row=7, end_column=mat_end
                    )
                # Apply matricula border to all cells in merge
                for c in range(mat_start, mat_end + 1):
                    ws.cell(row=7, column=c).border = border_matricula
            else:
                # MATERNAL: no separate SECCIONES area
                mat_start = grade_start
                mat_end = grade_end
                ws.cell(row=7, column=mat_start, value="MATRICULA").font = font_row7
                ws.cell(row=7, column=mat_start).alignment = align_center
                if mat_end > mat_start:
                    ws.merge_cells(
                        start_row=7, start_column=mat_start,
                        end_row=7, end_column=mat_end
                    )
                for c in range(mat_start, mat_end + 1):
                    ws.cell(row=7, column=c).border = border_matricula

        # ── Fixed columns A-G headers (rows 6-7) ──
        # A (col 1): N° - no header needed, just numbering
        # B (col 2): C.I.
        ws.cell(row=7, column=2, value="C.I.").font = font_row7
        ws.cell(row=7, column=2).alignment = align_center
        ws.cell(row=7, column=2).border = border_header_full
        # C (col 3): Nombre del Docente
        ws.cell(row=7, column=3, value="Nombre del Docente").font = font_row7
        ws.cell(row=7, column=3).alignment = align_center
        ws.cell(row=7, column=3).border = border_header_right

        # D (col 4): CARGA HORARIA
        ws.cell(row=6, column=4, value="CARGA HORARIA (RECIBO DE PAGO)").font = font_section_row6_small
        ws.cell(row=6, column=4).alignment = align_center_rot90
        ws.merge_cells(start_row=6, start_column=4, end_row=7, end_column=4)
        ws.cell(row=6, column=4).border = border_header_full

        # E (col 5): TITULAR/INTERINO
        ws.cell(row=6, column=5, value="TITULAR/ INTERINO").font = font_section_row6_small
        ws.cell(row=6, column=5).alignment = align_center_rot90
        ws.merge_cells(start_row=6, start_column=5, end_row=7, end_column=5)
        ws.cell(row=6, column=5).border = border_header_full

        # F (col 6): TURNO
        ws.cell(row=6, column=6, value="TURNO (MAÑANA, TARDE. INTEGRAL)").font = font_section_row6_small
        ws.cell(row=6, column=6).alignment = align_center_rot90
        ws.merge_cells(start_row=6, start_column=6, end_row=7, end_column=6)
        ws.cell(row=6, column=6).border = border_header_full

        # G (col 7): TIPO DE PERSONAL
        ws.cell(row=6, column=7, value="TIPO DE PERSONAL (Doc,Adm,Obr)").font = font_section_row6_small
        ws.cell(row=6, column=7).alignment = align_center_rot90
        ws.merge_cells(start_row=6, start_column=7, end_row=7, end_column=7)
        ws.cell(row=6, column=7).border = border_header_right

        # Apply borders to row 5 header cells
        for bk in BLOCK_ORDER:
            grade_start, grade_end = bloque_grade_cols[bk]
            for c in range(grade_start, grade_end + 1):
                cell = ws.cell(row=5, column=c)
                if c == grade_start:
                    cell.border = border_header_full
                elif c == grade_end and bloque_sec_cols[bk] is None:
                    cell.border = border_header_right
                else:
                    cell.border = border_header_full

            if bloque_sec_cols[bk] is not None:
                sec_start, sec_end = bloque_sec_cols[bk]
                for c in range(sec_start, sec_end + 1):
                    cell = ws.cell(row=5, column=c)
                    if c == sec_end:
                        cell.border = border_header_right
                    else:
                        cell.border = border_header_full

        # ── Row heights ──
        ws.row_dimensions[3].height = 21.75
        ws.row_dimensions[4].height = 16.5
        ws.row_dimensions[5].height = 17.25
        ws.row_dimensions[6].height = 144.75
        ws.row_dimensions[7].height = 17.25

        # ── Column widths ──
        ws.column_dimensions['A'].width = 4.4
        ws.column_dimensions['B'].width = 14.1
        ws.column_dimensions['C'].width = 33.7
        ws.column_dimensions['D'].width = 3.6
        ws.column_dimensions['E'].width = 4.7
        ws.column_dimensions['F'].width = 3.7
        ws.column_dimensions['G'].width = 3.7

        for bk in BLOCK_ORDER:
            grade_start, grade_end = bloque_grade_cols[bk]
            for c in range(grade_start, grade_end + 1):
                ws.column_dimensions[get_column_letter(c)].width = 3.7

            if bloque_sec_cols[bk] is not None:
                sec_start, sec_end = bloque_sec_cols[bk]
                # Gap column between grade and secciones
                ws.column_dimensions[get_column_letter(sec_start - 1)].width = 0.2
                for c in range(sec_start, sec_end + 1):
                    ws.column_dimensions[get_column_letter(c)].width = 3.7

        # ── Employee data rows ──
        data_start_row = 8

        # Build a mapping: (docente_id) -> {bloque_key: seccion_col}
        docente_seccion_map = {}
        for bk in BLOCK_ORDER:
            cfg = GRADO_MAP[bk]
            letras = bloque_letras[bk]
            grade_start, grade_end = bloque_grade_cols[bk]
            s_bd = bloque_secciones[bk]

            for s in s_bd:
                doc_id = s.get('docente_id')
                if doc_id is None:
                    continue
                letra = s.get('letra', '')
                col_idx = None
                if cfg['section_prefix'] == 'GRUPO':
                    try:
                        idx = int(letra) - 1 if letra.isdigit() else 0
                    except (ValueError, TypeError):
                        idx = 0
                    col_idx = grade_start + min(idx, grade_end - grade_start)
                else:
                    if letra in letras:
                        idx = letras.index(letra)
                        col_idx = grade_start + idx
                    elif letra:
                        try:
                            idx = ord(letra.upper()) - ord('A')
                            col_idx = grade_start + min(idx, grade_end - grade_start)
                        except (ValueError, TypeError):
                            continue
                if col_idx is not None:
                    if doc_id not in docente_seccion_map:
                        docente_seccion_map[doc_id] = {}
                    docente_seccion_map[doc_id][bk] = col_idx

        for idx, emp in enumerate(empleados, start=1):
            row = data_start_row + idx - 1
            ws.row_dimensions[row].height = 40

            cedula = str(emp.get('cedula', '')).strip()
            if cedula.startswith(('V-', 'E-', 'J-', 'G-')):
                cedula = cedula[2:]
            cedula = cedula.replace('.', '').replace('-', '')

            nombres = str(emp.get('nombres', '')).strip().upper()
            apellidos = str(emp.get('apellidos', '')).strip().upper()
            nombre_completo = f"{nombres} {apellidos}".strip()

            horas_acad = emp.get('horas_acad')
            horas_adm = emp.get('horas_adm')

            if horas_acad:
                carga_horaria = horas_acad
            elif horas_adm:
                carga_horaria = horas_adm
            else:
                carga_horaria = ''

            ws.cell(row=row, column=1, value=idx).font = font_data
            ws.cell(row=row, column=1).alignment = align_center
            ws.cell(row=row, column=1).border = border_data_first_last

            ws.cell(row=row, column=2, value=cedula).font = font_data
            ws.cell(row=row, column=2).alignment = align_center
            ws.cell(row=row, column=2).border = border_data_full

            ws.cell(row=row, column=3, value=nombre_completo).font = font_data
            ws.cell(row=row, column=3).alignment = align_left
            ws.cell(row=row, column=3).border = border_data_right

            ws.cell(row=row, column=4, value=carga_horaria).font = font_data
            ws.cell(row=row, column=4).alignment = align_center
            ws.cell(row=row, column=4).border = border_data_all_thin

            ws.cell(row=row, column=5, value="T").font = font_data
            ws.cell(row=row, column=5).alignment = align_center
            ws.cell(row=row, column=5).border = border_data_all_thin

            ws.cell(row=row, column=6, value="I").font = font_data
            ws.cell(row=row, column=6).alignment = align_center
            ws.cell(row=row, column=6).border = border_data_all_thin

            tipo = str(emp.get('tipo_personal', '')).strip()
            ws.cell(row=row, column=7, value=tipo).font = font_data
            ws.cell(row=row, column=7).alignment = align_center
            ws.cell(row=row, column=7).border = border_data_right

            # Mark the employee's assigned section with an X or leave matrícula blank
            emp_id = emp.get('id')
            emp_sections = docente_seccion_map.get(emp_id, {})
            for bk in BLOCK_ORDER:
                grade_start, grade_end = bloque_grade_cols[bk]
                for c in range(grade_start, grade_end + 1):
                    cell = ws.cell(row=row, column=c)
                    cell.border = border_data_all_thin
                    cell.font = font_data
                    cell.alignment = align_center

                if bloque_sec_cols[bk] is not None:
                    sec_start, sec_end = bloque_sec_cols[bk]
                    for c in range(sec_start, sec_end + 1):
                        cell = ws.cell(row=row, column=c)
                        cell.border = border_data_all_thin
                        cell.font = font_data
                        cell.alignment = align_center

                if bk in emp_sections:
                    target_col = emp_sections[bk]
                    ws.cell(row=row, column=target_col).value = "X"
                    ws.cell(row=row, column=target_col).font = font_data_bold

            # Light fill for data rows
            for c in range(1, last_data_col + 1):
                ws.cell(row=row, column=c).fill = light_fill

        wb.save(archivo)

        crear_msgbox(
            parent, "Exportación exitosa",
            f"Cuadratura generada correctamente.\n\n"
            f"Total de empleados: {len(empleados)}\n"
            f"Archivo: {os.path.basename(archivo)}",
            QMessageBox.Icon.Information
        ).exec()

        return archivo

    except Exception as e:
        crear_msgbox(
            parent, "Error",
            f"Error al generar cuadratura:\n{str(e)}",
            QMessageBox.Icon.Critical
        ).exec()
        return None

def generar_historial_notas_pdf(estudiante: dict, notas: list) -> str:
    """Genera un PDF con el historial de notas completo del estudiante."""
    
    # Validar datos
    campos_est = ["Nombres", "Apellidos", "Cédula"]
    valido, mensaje = validar_datos_exportacion(estudiante, campos_est)
    if not valido:
        raise ValueError(mensaje)
    
    # Normalizar datos
    estudiante["Nombres"] = str(estudiante["Nombres"]).strip().upper()
    estudiante["Apellidos"] = str(estudiante["Apellidos"]).strip().upper()
    
    # Normalizar cédula Escolar
    cedula = str(estudiante["Cédula"]).strip()
    if not cedula.upper().startswith("CE-"):
        cedula_escolar = f"CE-{cedula}"
    else:
        cedula_escolar = cedula.upper()
    
    # Crear carpeta
    carpeta = os.path.join(os.getcwd(), "exportados", "Historial academico")
    ok, msg = crear_carpeta_segura(carpeta)
    if not ok:
        raise IOError(msg)
    
    nombre_base = sanitizar_nombre_archivo(f"Historial_Notas_{estudiante['Cédula']}")
    nombre_archivo = os.path.join(carpeta, f"{nombre_base}.pdf")
    
    try:
        doc = SimpleDocTemplate(
            nombre_archivo,
            pagesize=letter,
            leftMargin=70,
            rightMargin=70,
            topMargin=180,
            bottomMargin=50
        )

        story = [Paragraph("HISTORIAL DE NOTAS ACADÉMICAS", styles["Title"]), Spacer(1, 16)]

        # Datos del estudiante
        datos_estudiante = f"""
        <b>Estudiante:</b> {estudiante['Nombres']} {estudiante['Apellidos']}<br/>
        <b>Cédula Escolar:</b> {cedula_escolar}<br/>
        """
        
        story.append(Paragraph(datos_estudiante, styles["Normal"]))
        story.append(Spacer(1, 20))
        
        if not notas:
            # Sin historial de notas
            texto_sin_datos = Paragraph(
                "<i>No se encontró historial de notas para este estudiante.</i>",
                justificado
            )
            story.append(texto_sin_datos)
        else:
            # Agrupar notas por año escolar
            notas_por_anio = {}
            for nota in notas:
                anio_key = f"{nota['año_inicio']}-{nota['año_inicio']+1}"
                if anio_key not in notas_por_anio:
                    notas_por_anio[anio_key] = []
                notas_por_anio[anio_key].append(nota)
            
            # Crear tabla para cada año escolar
            for anio_escolar in sorted(notas_por_anio.keys()):
                notas_anio = notas_por_anio[anio_escolar]
                
                # Encabezado de año
                story.append(Paragraph(
                    f"<b>Año Escolar: {anio_escolar}</b>",
                    ParagraphStyle('AnoEscolar', parent=styles['Normal'], fontSize=11,
                                 textColor=colors.HexColor('#2C3E50'), spaceAfter=8,
                                 spaceBefore=12)
                ))
                
                # Obtener información de grado y sección (del primer registro)
                if notas_anio:
                    nivel = notas_anio[0].get('nivel', '-')
                    grado = notas_anio[0].get('grado', '-')
                    letra = notas_anio[0].get('letra', '-')
                    
                    info_seccion = f"<i>Nivel: {nivel} | Grado: {grado} | Sección: {letra}</i>"
                    story.append(Paragraph(info_seccion, 
                        ParagraphStyle('InfoSeccion', parent=styles['Normal'], fontSize=9,
                                     textColor=colors.HexColor('#7F8C8D'), spaceAfter=10)))
                
                # Agrupar notas del año por Área de Aprendizaje
                notas_por_area = {}
                for nota in notas_anio:
                    area = nota.get('area_aprendizaje', 'Sin área')
                    if area not in notas_por_area:
                        notas_por_area[area] = []
                    notas_por_area[area].append(nota)
                
                # Crear tabla con áreas de aprendizaje como separadores
                datos_tabla = [[
                    Paragraph("<b>Materia</b>", centrado),
                    Paragraph("<b>Lapso 1</b>", centrado),
                    Paragraph("<b>Lapso 2</b>", centrado),
                    Paragraph("<b>Lapso 3</b>", centrado),
                    Paragraph("<b>Nota Final</b>", centrado)
                ]]
                
                # Track de filas de área y de notas para estilos
                filas_area = []  # Índices de filas que son encabezados de área
                filas_notas = []  # (índice_fila, nota_dict) para colorear
                fila_idx = 1  # Empezamos en 1 por el header
                
                for area_nombre in sorted(notas_por_area.keys()):
                    notas_area = notas_por_area[area_nombre]
                    
                    # Fila de encabezado del área de aprendizaje (span en toda la fila)
                    datos_tabla.append([
                        Paragraph(
                            f"<b>{area_nombre}</b>",
                            ParagraphStyle('AreaHeader', parent=styles['Normal'],
                                         fontSize=8, textColor=colors.white)
                        ),
                        "", "", "", ""
                    ])
                    filas_area.append(fila_idx)
                    fila_idx += 1
                    
                    # Materias del área
                    for nota in notas_area:
                        materia = str(nota.get('materia', '-'))
                        
                        # Formatear notas: priorizar literal si existe
                        lapso1 = "-"
                        if nota.get('lapso_1_lit'):
                            lapso1 = str(nota['lapso_1_lit'])
                        elif nota.get('lapso_1') is not None:
                            lapso1 = str(nota['lapso_1'])
                        
                        lapso2 = "-"
                        if nota.get('lapso_2_lit'):
                            lapso2 = str(nota['lapso_2_lit'])
                        elif nota.get('lapso_2') is not None:
                            lapso2 = str(nota['lapso_2'])
                        
                        lapso3 = "-"
                        if nota.get('lapso_3_lit'):
                            lapso3 = str(nota['lapso_3_lit'])
                        elif nota.get('lapso_3') is not None:
                            lapso3 = str(nota['lapso_3'])
                        
                        nota_final = "-"
                        if nota.get('nota_final_literal'):
                            nota_final = str(nota['nota_final_literal'])
                        elif nota.get('nota_final') is not None:
                            nota_final = str(nota['nota_final'])
                        
                        datos_tabla.append([
                            Paragraph(f"    {materia}", ParagraphStyle('Materia', parent=styles['Normal'], fontSize=9)),
                            Paragraph(lapso1, centrado),
                            Paragraph(lapso2, centrado),
                            Paragraph(lapso3, centrado),
                            Paragraph(nota_final, centrado)
                        ])
                        filas_notas.append((fila_idx, nota))
                        fila_idx += 1
                
                # Crear tabla con estilos
                tabla = Table(datos_tabla, colWidths=[150, 70, 70, 70, 80])
                
                estilos_tabla = [
                    # Encabezado
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E5894')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                    
                    # Contenido general
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    
                    # Bordes
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    
                    # Padding
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    
                    # Alineación de materia a la izquierda
                    ('ALIGN', (0, 1), (0, -1), 'LEFT'),
                ]
                
                # Estilos para filas de área de aprendizaje
                for fa in filas_area:
                    estilos_tabla.extend([
                        ('BACKGROUND', (0, fa), (-1, fa), colors.HexColor('#3498DB')),
                        ('TEXTCOLOR', (0, fa), (-1, fa), colors.white),
                        ('SPAN', (0, fa), (-1, fa)),
                        ('ALIGN', (0, fa), (0, fa), 'LEFT'),
                    ])
                
                tabla.setStyle(TableStyle(estilos_tabla))
                
                # Colorear celdas de nota final según aprobación
                for fi, nota in filas_notas:
                    aprobado = nota.get('aprobado')
                    if aprobado is not None:
                        if aprobado == 0 or aprobado == False:
                            tabla.setStyle(TableStyle([
                                ('BACKGROUND', (4, fi), (4, fi), colors.HexColor('#FADBD8'))
                            ]))
                        elif aprobado == 1 or aprobado == True:
                            tabla.setStyle(TableStyle([
                                ('BACKGROUND', (4, fi), (4, fi), colors.HexColor('#D5F4E6'))
                            ]))
                
                story.append(tabla)
                story.append(Spacer(1, 20))
            
            # Leyenda de colores
            story.append(Spacer(1, 10))
            leyenda_style = ParagraphStyle(
                'Leyenda',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#7F8C8D')
            )
            
            leyenda = """
            <b>Leyenda:</b> <font color="#117A65">Verde = Aprobado</font> | 
            <font color="#C0392B">Rojo = Reprobado</font> | 
            <font color="#34495E">-  = Sin calificación</font>
            """
            story.append(Paragraph(leyenda, leyenda_style))
        
        # Generar PDF
        doc.build(story, onFirstPage=encabezado_y_pie, onLaterPages=encabezado_y_pie)
        
        return nombre_archivo
        
    except Exception as e:
        raise IOError(f"Error generando historial de notas en PDF: {e}")